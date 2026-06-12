"""
Report Service - Professional PDF reports (CSRD/ESG design)
"""
from datetime import datetime
from typing import Optional, Dict, Any, List
from uuid import UUID
import io
import logging

logger = logging.getLogger(__name__)

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, KeepTogether
)
from reportlab.graphics.shapes import Drawing, Rect, String, Circle, Line, PolyLine
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.charts.lineplots import LinePlot
from reportlab.graphics.charts.barcharts import VerticalBarChart, HorizontalBarChart
from reportlab.graphics.widgets.markers import makeMarker
from sqlalchemy import select, extract
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.data_entry import DataEntry
from app.models.esg_score import ESGScore
from app.models.organization import Organization

# ─── Color Palette ────────────────────────────────────────────────────────────
C_NAVY   = colors.HexColor('#0f172a')
C_GREEN  = colors.HexColor('#059669')
C_GREEN2 = colors.HexColor('#d1fae5')
C_BLUE   = colors.HexColor('#2563eb')
C_BLUE2  = colors.HexColor('#dbeafe')
C_PURPLE = colors.HexColor('#7c3aed')
C_PURPLE2= colors.HexColor('#ede9fe')
C_GRAY   = colors.HexColor('#6b7280')
C_LGRAY  = colors.HexColor('#f3f4f6')
C_MGRAY  = colors.HexColor('#e5e7eb')
C_WHITE  = colors.white
C_AMBER  = colors.HexColor('#d97706')
C_AMBER2 = colors.HexColor('#fef3c7')
C_RED    = colors.HexColor('#dc2626')

# ESRS Section definitions
ESRS_SECTIONS = [
    # Environmental
    {'code': 'E1', 'label': 'Changement climatique',      'pillar': 'environmental', 'color': C_GREEN},
    {'code': 'E2', 'label': 'Pollution',                   'pillar': 'environmental', 'color': C_GREEN},
    {'code': 'E3', 'label': 'Eau et ressources marines',   'pillar': 'environmental', 'color': C_GREEN},
    {'code': 'E4', 'label': 'Biodiversité',                'pillar': 'environmental', 'color': C_GREEN},
    {'code': 'E5', 'label': 'Utilisation des ressources',  'pillar': 'environmental', 'color': C_GREEN},
    # Social
    {'code': 'S1', 'label': 'Effectifs propres',           'pillar': 'social', 'color': C_BLUE},
    {'code': 'S2', 'label': 'Travailleurs chaîne de valeur','pillar': 'social', 'color': C_BLUE},
    {'code': 'S3', 'label': 'Communautés affectées',       'pillar': 'social', 'color': C_BLUE},
    {'code': 'S4', 'label': 'Consommateurs',               'pillar': 'social', 'color': C_BLUE},
    # Governance
    {'code': 'G1', 'label': 'Conduite des affaires',       'pillar': 'governance', 'color': C_PURPLE},
]

PAGE_W, PAGE_H = A4


def _hex(c):
    """Return hex string from reportlab color."""
    return '#%02x%02x%02x' % (int(c.red*255), int(c.green*255), int(c.blue*255))


def _humanize_period(months: Optional[int]) -> str:
    """Convert a number of months to a human-friendly window label.

    Used to surface the data window of an ESG score in reports.
    """
    if not months or months <= 0:
        return '—'
    months = int(months)
    if months < 12:
        return f'{months} mois'
    years = months / 12.0
    if abs(years - round(years)) < 0.01:
        y = int(round(years))
        return f'{y} an' if y == 1 else f'{y} ans'
    return f'{months} mois'


def _make_header_footer(report_title: str, year: int):
    """Return onFirstPage / onLaterPages callbacks for SimpleDocTemplate."""
    w, h = PAGE_W, PAGE_H
    gen_date = datetime.now().strftime('%d/%m/%Y')

    def later_pages(canvas, doc):
        canvas.saveState()
        # Green top bar
        canvas.setFillColor(C_GREEN)
        canvas.rect(0, h - 12*mm, w, 12*mm, fill=1, stroke=0)
        # Logo
        canvas.setFillColor(C_WHITE)
        canvas.setFont('Helvetica-Bold', 10)
        canvas.drawString(20*mm, h - 8*mm, 'ESG Flow')
        # Title
        canvas.setFont('Helvetica', 8)
        canvas.drawCentredString(w / 2, h - 8*mm, report_title[:65])
        # Page number
        canvas.drawRightString(w - 20*mm, h - 8*mm, f'Page {doc.page}')
        # Footer line
        canvas.setStrokeColor(C_MGRAY)
        canvas.setLineWidth(0.5)
        canvas.line(20*mm, 15*mm, w - 20*mm, 15*mm)
        canvas.setFillColor(C_GRAY)
        canvas.setFont('Helvetica', 7)
        canvas.drawString(20*mm, 9*mm,
                          f'© {year} ESG Flow · Confidentiel · Généré le {gen_date}')
        canvas.drawRightString(w - 20*mm, 9*mm, 'CSRD · ESRS · GRI · TCFD')
        canvas.restoreState()

    def first_page(canvas, doc):
        pass  # Cover page has its own design (flowables)

    return first_page, later_pages


class ReportService:
    """Service for generating professional ESG reports"""

    REPORT_TYPES = {
        'executive': {'name': 'Rapport Exécutif', 'description': 'Vue d\'ensemble pour la direction'},
        'detailed':  {'name': 'Rapport Détaillé',  'description': 'Analyse complète'},
        'csrd':      {'name': 'Rapport CSRD',      'description': 'Conforme directive européenne'},
        'gri':       {'name': 'Rapport GRI',       'description': 'Standards GRI 2021'},
        'tcfd':      {'name': 'Rapport TCFD',      'description': 'Risques climatiques'},
        'sfdr':      {'name': 'Rapport SFDR',      'description': 'Produits financiers durables — données PAI & taxonomie UE'},
        'carbon':    {'name': 'Bilan Carbone ADEME', 'description': 'Émissions GES Scopes 1, 2, 3 selon méthode ADEME/GHG Protocol'},
        'dpef':      {'name': 'DPEF',              'description': 'Déclaration de performance extra-financière — Art. L.225-102-1'},
    }

    def __init__(self, db: AsyncSession):
        self.db = db

    async def generate_report(
        self,
        tenant_id: UUID,
        report_type: str,
        organization_id: Optional[UUID] = None,
        period: str = 'annual',
        year: Optional[int] = None,
        format: str = 'pdf'
    ) -> bytes:
        if report_type not in self.REPORT_TYPES:
            raise ValueError(f"Unknown report type: {report_type}")

        _year = year or datetime.now().year
        data = await self._collect_data(tenant_id, organization_id, _year)

        if format == 'pdf':
            return await self._generate_pdf(report_type, data, _year)
        elif format == 'excel':
            return self._generate_excel(report_type, data, _year)
        elif format == 'word':
            return self._generate_word(report_type, data, _year)
        elif format == 'json':
            import json as _json
            return _json.dumps({
                'report_type': report_type,
                'year': _year,
                'org_name': data.get('org_name'),
                'stats': data['stats'],
                'sections': {
                    code: [
                        {'metric': e['metric_name'], 'value': str(e['value']), 'unit': e['unit'], 'status': e['verification_status']}
                        for e in entries
                    ]
                    for code, entries in data['by_section'].items()
                },
            }, ensure_ascii=False, indent=2).encode('utf-8')
        else:
            raise ValueError(f"Unsupported format: {format}")

    # ─── Excel Export ─────────────────────────────────────────────────────────

    def _generate_excel(self, report_type: str, data: Dict[str, Any], year: int) -> bytes:
        """Generate a multi-sheet CSRD Excel report using openpyxl."""
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()
        stats = data['stats']
        org_name = data.get('org_name', 'Organisation')
        by_section = data.get('by_section', {})
        entries = data.get('entries', {})

        # ── Color palette ──
        GREEN  = '059669'; GREEN_L  = 'D1FAE5'
        BLUE   = '2563EB'; BLUE_L   = 'DBEAFE'
        PURPLE = '7C3AED'; PURPLE_L = 'EDE9FE'
        NAVY   = '0F172A'; GRAY     = '6B7280'
        AMBER  = 'D97706'; AMBER_L  = 'FEF3C7'
        RED    = 'DC2626'; WHITE    = 'FFFFFF'
        LGRAY  = 'F3F4F6'

        def _fill(hex_color):
            return PatternFill('solid', fgColor=hex_color)

        def _font(bold=False, color=NAVY, size=10):
            return Font(bold=bold, color=color, size=size, name='Calibri')

        def _border():
            thin = Side(style='thin', color='E5E7EB')
            return Border(left=thin, right=thin, top=thin, bottom=thin)

        def _header_row(ws, row, values, fill_color, font_color=WHITE, bold=True):
            for col, val in enumerate(values, 1):
                c = ws.cell(row=row, column=col, value=val)
                c.fill = _fill(fill_color)
                c.font = Font(bold=bold, color=font_color, size=10, name='Calibri')
                c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                c.border = _border()

        # ══════════════════════════════════════════
        # Sheet 1 — Résumé exécutif
        # ══════════════════════════════════════════
        ws_summary = wb.active
        ws_summary.title = 'Résumé exécutif'
        ws_summary.sheet_view.showGridLines = False
        ws_summary.column_dimensions['A'].width = 30
        ws_summary.column_dimensions['B'].width = 25
        ws_summary.column_dimensions['C'].width = 20

        # Title block
        ws_summary.merge_cells('A1:C1')
        c = ws_summary['A1']
        c.value = f'Rapport CSRD — {org_name} — {year}'
        c.font = Font(bold=True, color=WHITE, size=14, name='Calibri')
        c.fill = _fill(NAVY)
        c.alignment = Alignment(horizontal='center', vertical='center')
        ws_summary.row_dimensions[1].height = 36

        ws_summary.merge_cells('A2:C2')
        c = ws_summary['A2']
        c.value = f'Conforme ESRS 2024 — Généré le {datetime.now().strftime("%d/%m/%Y")} — ESG Flow'
        c.font = Font(color=GRAY, size=9, name='Calibri')
        c.fill = _fill(LGRAY)
        c.alignment = Alignment(horizontal='center')
        ws_summary.row_dimensions[2].height = 20

        kpis = [
            ('Score ESG Global', f'{stats["score"]}/100', GREEN),
            ('Indicateurs collectés', str(stats['total_entries']), BLUE),
            ('Vérifiés', str(stats['verified_count']), GREEN),
            ('En attente', str(stats['pending_count']), AMBER),
            ('Couverture Environnement', f'{stats["by_pillar"]["environmental"]} indicateurs', GREEN),
            ('Couverture Social', f'{stats["by_pillar"]["social"]} indicateurs', BLUE),
            ('Couverture Gouvernance', f'{stats["by_pillar"]["governance"]} indicateurs', PURPLE),
        ]

        ws_summary.row_dimensions[3].height = 12
        for i, (label, value, color) in enumerate(kpis, start=4):
            ws_summary.cell(row=i, column=1, value=label).font = _font(bold=True)
            ws_summary.cell(row=i, column=1).fill = _fill(LGRAY)
            ws_summary.cell(row=i, column=1).border = _border()
            c_val = ws_summary.cell(row=i, column=2, value=value)
            c_val.font = Font(bold=True, color=color, size=12, name='Calibri')
            c_val.alignment = Alignment(horizontal='center')
            c_val.border = _border()
            ws_summary.merge_cells(f'C{i}:C{i}')
            ws_summary.row_dimensions[i].height = 22

        # ══════════════════════════════════════════
        # Sheet 2 — Données ESRS (all sections)
        # ══════════════════════════════════════════
        ws_esrs = wb.create_sheet('Données ESRS')
        ws_esrs.sheet_view.showGridLines = False
        ws_esrs.freeze_panes = 'A2'
        ws_esrs.column_dimensions['A'].width = 8
        ws_esrs.column_dimensions['B'].width = 35
        ws_esrs.column_dimensions['C'].width = 18
        ws_esrs.column_dimensions['D'].width = 12
        ws_esrs.column_dimensions['E'].width = 12
        ws_esrs.column_dimensions['F'].width = 20

        _header_row(ws_esrs, 1, ['Section', 'Indicateur', 'Valeur', 'Unité', 'Catégorie', 'Statut'], NAVY)
        ws_esrs.row_dimensions[1].height = 22

        pillar_colors = {
            'environmental': (GREEN, GREEN_L),
            'social':        (BLUE,  BLUE_L),
            'governance':    (PURPLE, PURPLE_L),
        }
        section_pillar_map = {s['code']: s['pillar'] for s in ESRS_SECTIONS}

        row = 2
        for sec in ESRS_SECTIONS:
            sec_entries = by_section.get(sec['code'], [])
            pillar = section_pillar_map.get(sec['code'], 'environmental')
            main_color, light_color = pillar_colors.get(pillar, (GREEN, GREEN_L))

            if not sec_entries:
                ws_esrs.cell(row=row, column=1, value=sec['code']).fill = _fill(AMBER_L)
                ws_esrs.cell(row=row, column=1).font = _font(bold=True, color=AMBER)
                ws_esrs.cell(row=row, column=1).border = _border()
                c = ws_esrs.cell(row=row, column=2, value='Aucune donnée collectée')
                c.fill = _fill(AMBER_L)
                c.font = Font(italic=True, color=AMBER, size=9, name='Calibri')
                c.border = _border()
                for col in range(3, 7):
                    ws_esrs.cell(row=row, column=col).fill = _fill(AMBER_L)
                    ws_esrs.cell(row=row, column=col).border = _border()
                row += 1
                continue

            for j, entry in enumerate(sec_entries):
                bg = light_color if j % 2 == 0 else WHITE
                ws_esrs.cell(row=row, column=1, value=sec['code'] if j == 0 else '').fill = _fill(main_color if j == 0 else bg)
                ws_esrs.cell(row=row, column=1).font = _font(bold=(j == 0), color=WHITE if j == 0 else NAVY)
                ws_esrs.cell(row=row, column=1).border = _border()
                ws_esrs.cell(row=row, column=2, value=str(entry.get('metric_name', ''))[:60]).fill = _fill(bg)
                ws_esrs.cell(row=row, column=2).border = _border()
                ws_esrs.cell(row=row, column=3, value=str(entry.get('value', '—'))).fill = _fill(bg)
                ws_esrs.cell(row=row, column=3).alignment = Alignment(horizontal='center')
                ws_esrs.cell(row=row, column=3).border = _border()
                ws_esrs.cell(row=row, column=4, value=str(entry.get('unit', ''))).fill = _fill(bg)
                ws_esrs.cell(row=row, column=4).alignment = Alignment(horizontal='center')
                ws_esrs.cell(row=row, column=4).border = _border()
                ws_esrs.cell(row=row, column=5, value=str(entry.get('category', ''))).fill = _fill(bg)
                ws_esrs.cell(row=row, column=5).border = _border()
                status = entry.get('verification_status', 'pending')
                status_label = {'verified': '✓ Vérifié', 'pending': '⏳ En attente', 'rejected': '✗ Rejeté'}.get(status, status)
                status_color = GREEN if status == 'verified' else (AMBER if status == 'pending' else RED)
                ws_esrs.cell(row=row, column=6, value=status_label).fill = _fill(bg)
                ws_esrs.cell(row=row, column=6).font = Font(color=status_color, size=9, name='Calibri')
                ws_esrs.cell(row=row, column=6).alignment = Alignment(horizontal='center')
                ws_esrs.cell(row=row, column=6).border = _border()
                ws_esrs.row_dimensions[row].height = 18
                row += 1

        # ══════════════════════════════════════════
        # Sheet 3 — Gap Analysis
        # ══════════════════════════════════════════
        ws_gap = wb.create_sheet('Gap Analysis')
        ws_gap.sheet_view.showGridLines = False
        ws_gap.freeze_panes = 'A2'
        ws_gap.column_dimensions['A'].width = 10
        ws_gap.column_dimensions['B'].width = 28
        ws_gap.column_dimensions['C'].width = 12
        ws_gap.column_dimensions['D'].width = 14
        ws_gap.column_dimensions['E'].width = 16
        ws_gap.column_dimensions['F'].width = 32

        _header_row(ws_gap, 1, ['Code', 'Section ESRS', 'Indicateurs', 'Couverture %', 'Statut', 'Action recommandée'], NAVY)
        ws_gap.row_dimensions[1].height = 22

        for i, sec in enumerate(ESRS_SECTIONS):
            row_num = i + 2
            count = len(by_section.get(sec['code'], []))
            pct = min(round(count / 5 * 100), 100) if count else 0

            if pct >= 80:
                status, action, row_color = '✓ Prêt', 'Maintenir & vérifier', GREEN_L
                st_color = GREEN
            elif pct >= 30:
                status, action, row_color = '⚠ Partiel', 'Compléter les données', AMBER_L
                st_color = AMBER
            else:
                status, action, row_color = '✗ Lacune', 'Collecte de données urgente', 'FEE2E2'
                st_color = RED

            bg = _fill(row_color)
            ws_gap.cell(row=row_num, column=1, value=sec['code']).fill = bg
            ws_gap.cell(row=row_num, column=1).font = _font(bold=True)
            ws_gap.cell(row=row_num, column=1).border = _border()
            ws_gap.cell(row=row_num, column=2, value=sec['label']).fill = bg
            ws_gap.cell(row=row_num, column=2).border = _border()
            ws_gap.cell(row=row_num, column=3, value=count).fill = bg
            ws_gap.cell(row=row_num, column=3).alignment = Alignment(horizontal='center')
            ws_gap.cell(row=row_num, column=3).border = _border()
            ws_gap.cell(row=row_num, column=4, value=f'{pct}%').fill = bg
            ws_gap.cell(row=row_num, column=4).alignment = Alignment(horizontal='center')
            ws_gap.cell(row=row_num, column=4).font = Font(bold=True, color=st_color, size=10, name='Calibri')
            ws_gap.cell(row=row_num, column=4).border = _border()
            ws_gap.cell(row=row_num, column=5, value=status).fill = bg
            ws_gap.cell(row=row_num, column=5).font = Font(bold=True, color=st_color, size=10, name='Calibri')
            ws_gap.cell(row=row_num, column=5).alignment = Alignment(horizontal='center')
            ws_gap.cell(row=row_num, column=5).border = _border()
            ws_gap.cell(row=row_num, column=6, value=action).fill = bg
            ws_gap.cell(row=row_num, column=6).border = _border()
            ws_gap.row_dimensions[row_num].height = 20

        # ══════════════════════════════════════════
        # Sheet 4 — Par pilier
        # ══════════════════════════════════════════
        for pillar_key, pillar_label, pcolor, plcolor in [
            ('environmental', 'Environnement', GREEN, GREEN_L),
            ('social',        'Social',        BLUE,  BLUE_L),
            ('governance',    'Gouvernance',   PURPLE, PURPLE_L),
        ]:
            ws_p = wb.create_sheet(pillar_label[:15])
            ws_p.sheet_view.showGridLines = False
            ws_p.freeze_panes = 'A2'
            ws_p.column_dimensions['A'].width = 35
            ws_p.column_dimensions['B'].width = 18
            ws_p.column_dimensions['C'].width = 12
            ws_p.column_dimensions['D'].width = 15

            _header_row(ws_p, 1, ['Métrique', 'Valeur', 'Unité', 'Statut'], pcolor)
            ws_p.row_dimensions[1].height = 22

            pillar_entries = entries.get(pillar_key, [])
            for j, entry in enumerate(pillar_entries):
                rn = j + 2
                bg = _fill(plcolor if j % 2 == 0 else WHITE)
                ws_p.cell(row=rn, column=1, value=str(entry.get('metric_name', ''))[:60]).fill = bg
                ws_p.cell(row=rn, column=1).border = _border()
                ws_p.cell(row=rn, column=2, value=str(entry.get('value', '—'))).fill = bg
                ws_p.cell(row=rn, column=2).alignment = Alignment(horizontal='center')
                ws_p.cell(row=rn, column=2).border = _border()
                ws_p.cell(row=rn, column=3, value=str(entry.get('unit', ''))).fill = bg
                ws_p.cell(row=rn, column=3).alignment = Alignment(horizontal='center')
                ws_p.cell(row=rn, column=3).border = _border()
                status = entry.get('verification_status', 'pending')
                status_label = {'verified': '✓', 'pending': '⏳', 'rejected': '✗'}.get(status, '?')
                s_color = GREEN if status == 'verified' else (AMBER if status == 'pending' else RED)
                ws_p.cell(row=rn, column=4, value=status_label).fill = bg
                ws_p.cell(row=rn, column=4).font = Font(color=s_color, size=11, name='Calibri')
                ws_p.cell(row=rn, column=4).alignment = Alignment(horizontal='center')
                ws_p.cell(row=rn, column=4).border = _border()
                ws_p.row_dimensions[rn].height = 18

            if not pillar_entries:
                ws_p.cell(row=2, column=1, value='Aucune donnée pour ce pilier').font = Font(italic=True, color=GRAY, name='Calibri')

        # ══════════════════════════════════════════
        # Sheet — Historique des scores (12 derniers points)
        # ══════════════════════════════════════════
        score_history = data.get('score_history', [])
        if score_history:
            ws_h = wb.create_sheet('Historique scores')
            ws_h.sheet_view.showGridLines = False
            ws_h.column_dimensions['A'].width = 14
            for col_letter in ('B', 'C', 'D', 'E', 'F', 'G'):
                ws_h.column_dimensions[col_letter].width = 16

            ws_h.merge_cells('A1:G1')
            t = ws_h['A1']
            t.value = f'Historique du score ESG — {org_name}'
            t.font = Font(bold=True, color=WHITE, size=13, name='Calibri')
            t.fill = _fill(NAVY)
            t.alignment = Alignment(horizontal='center', vertical='center')
            ws_h.row_dimensions[1].height = 30

            _header_row(ws_h, 3, ['Date', 'Score global', 'Note', 'Environnement', 'Social', 'Gouvernance', 'Complétude (%)'], NAVY)
            for i, s in enumerate(score_history, start=4):
                ws_h.cell(row=i, column=1, value=s.get('date') or '—').alignment = Alignment(horizontal='center')
                ws_h.cell(row=i, column=2, value=s.get('overall')).alignment = Alignment(horizontal='center')
                ws_h.cell(row=i, column=3, value=s.get('rating')).alignment = Alignment(horizontal='center')
                ws_h.cell(row=i, column=4, value=s.get('environmental')).alignment = Alignment(horizontal='center')
                ws_h.cell(row=i, column=5, value=s.get('social')).alignment = Alignment(horizontal='center')
                ws_h.cell(row=i, column=6, value=s.get('governance')).alignment = Alignment(horizontal='center')
                ws_h.cell(row=i, column=7, value=s.get('data_completeness')).alignment = Alignment(horizontal='center')
                for c in range(1, 8):
                    ws_h.cell(row=i, column=c).border = _border()

        # ══════════════════════════════════════════
        # Sheet — Matrice de matérialité (double matérialité)
        # ══════════════════════════════════════════
        materiality_issues = data.get('materiality_issues', [])
        if materiality_issues:
            ws_m = wb.create_sheet('Matérialité')
            ws_m.sheet_view.showGridLines = False
            widths_m = {'A': 30, 'B': 18, 'C': 16, 'D': 16, 'E': 12, 'F': 12, 'G': 30}
            for col, w in widths_m.items():
                ws_m.column_dimensions[col].width = w

            ws_m.merge_cells('A1:G1')
            t = ws_m['A1']
            t.value = f'Analyse de double matérialité — {org_name} — {year}'
            t.font = Font(bold=True, color=WHITE, size=13, name='Calibri')
            t.fill = _fill(PURPLE)
            t.alignment = Alignment(horizontal='center', vertical='center')
            ws_m.row_dimensions[1].height = 30

            _header_row(ws_m, 3, ['Enjeu', 'Catégorie', 'Impact financier', 'Impact ESG', 'Matériel ?', 'Priorité', 'Parties prenantes'], PURPLE)
            for i, m in enumerate(materiality_issues, start=4):
                ws_m.cell(row=i, column=1, value=m.get('name'))
                ws_m.cell(row=i, column=2, value=m.get('category'))
                ws_m.cell(row=i, column=3, value=m.get('financial_impact'))
                ws_m.cell(row=i, column=4, value=m.get('esg_impact'))
                ws_m.cell(row=i, column=5, value='✓' if m.get('is_material') else '—')
                ws_m.cell(row=i, column=6, value=m.get('priority'))
                ws_m.cell(row=i, column=7, value=m.get('stakeholders') or '—')
                for c in range(1, 8):
                    ws_m.cell(row=i, column=c).border = _border()
                    ws_m.cell(row=i, column=c).alignment = Alignment(vertical='center', wrap_text=True)
                ws_m.cell(row=i, column=5).alignment = Alignment(horizontal='center', vertical='center')

        # ══════════════════════════════════════════
        # Sheet — Registre des risques ESG
        # ══════════════════════════════════════════
        esg_risks = data.get('esg_risks', [])
        if esg_risks:
            ws_r = wb.create_sheet('Risques ESG')
            ws_r.sheet_view.showGridLines = False
            widths_r = {'A': 30, 'B': 18, 'C': 12, 'D': 10, 'E': 12, 'F': 14, 'G': 18, 'H': 24}
            for col, w in widths_r.items():
                ws_r.column_dimensions[col].width = w

            ws_r.merge_cells('A1:H1')
            t = ws_r['A1']
            t.value = f'Registre des risques ESG — {org_name} — {year}'
            t.font = Font(bold=True, color=WHITE, size=13, name='Calibri')
            t.fill = _fill(RED)
            t.alignment = Alignment(horizontal='center', vertical='center')
            ws_r.row_dimensions[1].height = 30

            _header_row(ws_r, 3, ['Risque', 'Catégorie', 'Probabilité', 'Impact', 'Score', 'Sévérité', 'Mitigation', 'Responsable'], RED)
            for i, r in enumerate(esg_risks, start=4):
                ws_r.cell(row=i, column=1, value=r.get('title'))
                ws_r.cell(row=i, column=2, value=r.get('category'))
                ws_r.cell(row=i, column=3, value=r.get('probability'))
                ws_r.cell(row=i, column=4, value=r.get('impact'))
                score = r.get('risk_score') or (int(r.get('probability', 0)) * int(r.get('impact', 0)))
                cell = ws_r.cell(row=i, column=5, value=score)
                cell.font = Font(
                    bold=True,
                    color=RED if score >= 16 else (AMBER if score >= 9 else GREEN),
                    size=10,
                    name='Calibri',
                )
                ws_r.cell(row=i, column=6, value=r.get('severity'))
                ws_r.cell(row=i, column=7, value=r.get('mitigation_status'))
                ws_r.cell(row=i, column=8, value=r.get('responsible') or '—')
                for c in range(1, 9):
                    ws_r.cell(row=i, column=c).border = _border()
                    ws_r.cell(row=i, column=c).alignment = Alignment(vertical='center', wrap_text=True)
                for c in (3, 4, 5):
                    ws_r.cell(row=i, column=c).alignment = Alignment(horizontal='center', vertical='center')

        # ══════════════════════════════════════════
        # Sheet — Fournisseurs (Supply chain ESG)
        # ══════════════════════════════════════════
        suppliers = data.get('suppliers', [])
        if suppliers:
            ws_s = wb.create_sheet('Fournisseurs')
            ws_s.sheet_view.showGridLines = False
            widths_s = {'A': 28, 'B': 14, 'C': 16, 'D': 12, 'E': 12, 'F': 14, 'G': 12, 'H': 12, 'I': 12, 'J': 12}
            for col, w in widths_s.items():
                ws_s.column_dimensions[col].width = w

            ws_s.merge_cells('A1:J1')
            t = ws_s['A1']
            t.value = f'Chaîne d\'approvisionnement ESG — {org_name}'
            t.font = Font(bold=True, color=WHITE, size=13, name='Calibri')
            t.fill = _fill(BLUE)
            t.alignment = Alignment(horizontal='center', vertical='center')
            ws_s.row_dimensions[1].height = 30

            _header_row(ws_s, 3, ['Nom', 'Pays', 'Catégorie', 'Risque', 'Statut', 'Dépense (k€)', 'Score global', 'E', 'S', 'G'], BLUE)
            for i, sup in enumerate(suppliers, start=4):
                ws_s.cell(row=i, column=1, value=sup.get('name'))
                ws_s.cell(row=i, column=2, value=sup.get('country'))
                ws_s.cell(row=i, column=3, value=sup.get('category'))
                ws_s.cell(row=i, column=4, value=sup.get('risk_level'))
                ws_s.cell(row=i, column=5, value=sup.get('status'))
                ws_s.cell(row=i, column=6, value=sup.get('spend_k_eur'))
                ws_s.cell(row=i, column=7, value=sup.get('global_score'))
                ws_s.cell(row=i, column=8, value=sup.get('env_score'))
                ws_s.cell(row=i, column=9, value=sup.get('social_score'))
                ws_s.cell(row=i, column=10, value=sup.get('gov_score'))
                for c in range(1, 11):
                    ws_s.cell(row=i, column=c).border = _border()
                for c in (4, 5, 6, 7, 8, 9, 10):
                    ws_s.cell(row=i, column=c).alignment = Alignment(horizontal='center', vertical='center')

        # ══════════════════════════════════════════
        # Sheet — Catalogue d'indicateurs (indicator_data)
        # ══════════════════════════════════════════
        indicator_catalog = data.get('indicator_catalog', [])
        if indicator_catalog:
            ws_i = wb.create_sheet('Indicateurs (catalogue)')
            ws_i.sheet_view.showGridLines = False
            widths_i = {'A': 12, 'B': 12, 'C': 32, 'D': 14, 'E': 10, 'F': 14, 'G': 12, 'H': 14}
            for col, w in widths_i.items():
                ws_i.column_dimensions[col].width = w

            ws_i.merge_cells('A1:H1')
            t = ws_i['A1']
            t.value = f'Catalogue d\'indicateurs renseignés — {len(indicator_catalog)} entrées'
            t.font = Font(bold=True, color=WHITE, size=13, name='Calibri')
            t.fill = _fill(GREEN)
            t.alignment = Alignment(horizontal='center', vertical='center')
            ws_i.row_dimensions[1].height = 30

            _header_row(ws_i, 3, ['Date', 'Code', 'Indicateur', 'Valeur', 'Unité', 'Source', 'Vérifié', 'Statut'], GREEN)
            for i, ic in enumerate(indicator_catalog, start=4):
                ws_i.cell(row=i, column=1, value=ic.get('date'))
                ws_i.cell(row=i, column=2, value=ic.get('code'))
                ws_i.cell(row=i, column=3, value=ic.get('name'))
                ws_i.cell(row=i, column=4, value=ic.get('value'))
                ws_i.cell(row=i, column=5, value=ic.get('unit'))
                ws_i.cell(row=i, column=6, value=ic.get('source'))
                ws_i.cell(row=i, column=7, value='✓' if ic.get('is_verified') else '—')
                ws_i.cell(row=i, column=8, value=ic.get('validation_status'))
                for c in range(1, 9):
                    ws_i.cell(row=i, column=c).border = _border()
                ws_i.cell(row=i, column=7).alignment = Alignment(horizontal='center')

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.read()

    # ─── Word Export ──────────────────────────────────────────────────────────

    def _generate_word(self, report_type: str, data: Dict[str, Any], year: int) -> bytes:
        """Generate a structured Word (.docx) CSRD report."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise ValueError("python-docx library not installed. Please add python-docx to requirements.txt.")

        stats = data['stats']
        org_name = data.get('org_name', 'Organisation')
        by_section = data.get('by_section', {})
        entries = data.get('entries', {})

        doc = Document()

        # Page setup: A4
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)

        # ── Cover Page ──
        doc.add_heading('RAPPORT CSRD', level=0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(org_name)
        run.bold = True
        run.font.size = Pt(18)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f'Exercice {year} · Conforme ESRS 2024').font.size = Pt(12)

        doc.add_paragraph(f'Score ESG Global : {stats["score"]}/100')
        doc.add_paragraph(f'Notation : {stats["rating"]}')
        doc.add_paragraph(f'Indicateurs collectés : {stats["total_entries"]} (vérifiés : {stats["verified_count"]})')
        doc.add_paragraph(f'Généré le : {datetime.now().strftime("%d/%m/%Y à %H:%M")} — ESG Flow')

        doc.add_page_break()

        # ── Executive Summary ──
        doc.add_heading('1. Résumé Exécutif', level=1)
        summary_data = [
            ['Indicateur', 'Valeur'],
            ['Score ESG Global', f'{stats["score"]}/100'],
            ['Indicateurs totaux', str(stats['total_entries'])],
            ['Indicateurs vérifiés', str(stats['verified_count'])],
            ['En attente de vérification', str(stats['pending_count'])],
            ['Indicateurs Environnement', str(stats['by_pillar']['environmental'])],
            ['Indicateurs Social', str(stats['by_pillar']['social'])],
            ['Indicateurs Gouvernance', str(stats['by_pillar']['governance'])],
        ]
        t = doc.add_table(rows=len(summary_data), cols=2)
        t.style = 'Table Grid'
        for i, row_data in enumerate(summary_data):
            for j, cell_val in enumerate(row_data):
                cell = t.cell(i, j)
                cell.text = cell_val
                if i == 0:
                    cell.paragraphs[0].runs[0].bold = True

        doc.add_page_break()

        # ── ESRS Sections ──
        doc.add_heading('2. Données ESRS par Standard', level=1)

        for sec in ESRS_SECTIONS:
            sec_entries = by_section.get(sec['code'], [])
            count = len(sec_entries)
            pct = min(round(count / 5 * 100), 100) if count else 0

            doc.add_heading(f'{sec["code"]} — {sec["label"]} ({pct}% couvert)', level=2)

            if not sec_entries:
                p = doc.add_paragraph('⚠ Aucune donnée collectée pour cette section.')
                p.runs[0].italic = True
                continue

            headers = ['Métrique', 'Valeur', 'Unité', 'Statut']
            t = doc.add_table(rows=1 + len(sec_entries[:20]), cols=len(headers))
            t.style = 'Table Grid'

            hdr_row = t.rows[0]
            for j, h in enumerate(headers):
                hdr_row.cells[j].text = h
                hdr_row.cells[j].paragraphs[0].runs[0].bold = True

            for i, entry in enumerate(sec_entries[:20]):
                row = t.rows[i + 1]
                row.cells[0].text = str(entry.get('metric_name', ''))[:60]
                row.cells[1].text = str(entry.get('value', '—'))
                row.cells[2].text = str(entry.get('unit', ''))
                status = entry.get('verification_status', 'pending')
                row.cells[3].text = {'verified': '✓ Vérifié', 'pending': '⏳ En attente', 'rejected': '✗ Rejeté'}.get(status, status)

        doc.add_page_break()

        # ── Gap Analysis ──
        doc.add_heading('3. Analyse des Écarts (Gap Analysis)', level=1)
        gap_headers = ['Code', 'Section', 'Indicateurs', 'Couverture', 'Statut', 'Action']
        gt = doc.add_table(rows=1 + len(ESRS_SECTIONS), cols=len(gap_headers))
        gt.style = 'Table Grid'

        hdr_row = gt.rows[0]
        for j, h in enumerate(gap_headers):
            hdr_row.cells[j].text = h
            hdr_row.cells[j].paragraphs[0].runs[0].bold = True

        for i, sec in enumerate(ESRS_SECTIONS):
            count = len(by_section.get(sec['code'], []))
            pct = min(round(count / 5 * 100), 100) if count else 0
            row = gt.rows[i + 1]
            row.cells[0].text = sec['code']
            row.cells[1].text = sec['label']
            row.cells[2].text = str(count)
            row.cells[3].text = f'{pct}%'
            if pct >= 80:
                row.cells[4].text = '✓ Prêt'
                row.cells[5].text = 'Maintenir & vérifier'
            elif pct >= 30:
                row.cells[4].text = '⚠ Partiel'
                row.cells[5].text = 'Compléter les données'
            else:
                row.cells[4].text = '✗ Lacune'
                row.cells[5].text = 'Collecte urgente'

        doc.add_page_break()

        # ── Footer info ──
        doc.add_heading('4. Cadres de Référence', level=1)
        frameworks = [
            'CSRD — Corporate Sustainability Reporting Directive (UE) 2022/2464',
            'ESRS — European Sustainability Reporting Standards 2024',
            'GRI — Global Reporting Initiative Standards 2021',
            'TCFD — Task Force on Climate-related Financial Disclosures',
            'SFDR — Sustainable Finance Disclosure Regulation',
            'Taxonomie UE — Règlement (UE) 2020/852',
        ]
        for fw in frameworks:
            doc.add_paragraph(fw, style='List Bullet')

        # ── Annexes : matérialité, risques, fournisseurs, historique scores ──
        materiality_issues = data.get('materiality_issues', [])
        if materiality_issues:
            doc.add_page_break()
            doc.add_heading('Annexe A — Matrice de double matérialité', level=1)
            doc.add_paragraph(
                f"{len(materiality_issues)} enjeu(x) ESG identifié(s). Les enjeux marqués comme 'matériels' "
                "satisfont au critère ESRS de double matérialité (impact financier ET impact ESG significatifs)."
            )
            mat_t = doc.add_table(rows=1 + min(len(materiality_issues), 30), cols=5)
            mat_t.style = 'Table Grid'
            for j, h in enumerate(['Enjeu', 'Catégorie', 'Imp. financier', 'Imp. ESG', 'Matériel ?']):
                mat_t.rows[0].cells[j].text = h
                mat_t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            for i, m in enumerate(materiality_issues[:30]):
                row = mat_t.rows[i + 1]
                row.cells[0].text = m.get('name') or '—'
                row.cells[1].text = m.get('category') or '—'
                row.cells[2].text = f"{m.get('financial_impact', 0):.0f}/100"
                row.cells[3].text = f"{m.get('esg_impact', 0):.0f}/100"
                row.cells[4].text = '✓' if m.get('is_material') else '—'

        esg_risks = data.get('esg_risks', [])
        if esg_risks:
            doc.add_page_break()
            doc.add_heading('Annexe B — Registre des risques ESG', level=1)
            doc.add_paragraph(
                f"{len(esg_risks)} risque(s) ESG identifié(s) et évalué(s) selon une matrice probabilité × impact."
            )
            risk_t = doc.add_table(rows=1 + min(len(esg_risks), 30), cols=5)
            risk_t.style = 'Table Grid'
            for j, h in enumerate(['Risque', 'Catégorie', 'Score', 'Sévérité', 'Mitigation']):
                risk_t.rows[0].cells[j].text = h
                risk_t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            for i, r in enumerate(esg_risks[:30]):
                row = risk_t.rows[i + 1]
                row.cells[0].text = r.get('title') or '—'
                row.cells[1].text = r.get('category') or '—'
                score = r.get('risk_score') or (int(r.get('probability', 0)) * int(r.get('impact', 0)))
                row.cells[2].text = str(score)
                row.cells[3].text = r.get('severity') or '—'
                row.cells[4].text = r.get('mitigation_status') or '—'

        suppliers = data.get('suppliers', [])
        if suppliers:
            doc.add_page_break()
            doc.add_heading('Annexe C — Fournisseurs notables (top 20)', level=1)
            doc.add_paragraph(
                f"Évaluation ESG de {min(len(suppliers), 20)} fournisseurs principaux sur un total de {len(suppliers)}. "
                "Les scores reflètent la performance ESG auto-déclarée et auditée."
            )
            sup_t = doc.add_table(rows=1 + min(len(suppliers), 20), cols=5)
            sup_t.style = 'Table Grid'
            for j, h in enumerate(['Fournisseur', 'Pays', 'Risque', 'Statut', 'Score global']):
                sup_t.rows[0].cells[j].text = h
                sup_t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            for i, sup in enumerate(suppliers[:20]):
                row = sup_t.rows[i + 1]
                row.cells[0].text = sup.get('name') or '—'
                row.cells[1].text = sup.get('country') or '—'
                row.cells[2].text = sup.get('risk_level') or '—'
                row.cells[3].text = sup.get('status') or '—'
                gs = sup.get('global_score')
                row.cells[4].text = f"{gs:.1f}/100" if gs is not None else '—'

        score_history = data.get('score_history', [])
        if score_history:
            doc.add_page_break()
            doc.add_heading('Annexe D — Historique du score ESG', level=1)
            doc.add_paragraph(
                f"Évolution du score ESG global et par pilier sur les {len(score_history)} dernières mesures."
            )
            sh_t = doc.add_table(rows=1 + len(score_history), cols=6)
            sh_t.style = 'Table Grid'
            for j, h in enumerate(['Date', 'Global', 'Note', 'Env.', 'Soc.', 'Gouv.']):
                sh_t.rows[0].cells[j].text = h
                sh_t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            for i, s in enumerate(score_history):
                row = sh_t.rows[i + 1]
                row.cells[0].text = s.get('date') or '—'
                row.cells[1].text = str(s.get('overall', '—'))
                row.cells[2].text = s.get('rating') or '—'
                row.cells[3].text = str(s.get('environmental', '—'))
                row.cells[4].text = str(s.get('social', '—'))
                row.cells[5].text = str(s.get('governance', '—'))

        p = doc.add_paragraph()
        p.add_run(f'\n© {year} ESG Flow · Rapport généré automatiquement · {datetime.now().strftime("%d/%m/%Y")}').italic = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # Mapping pillar+keywords → ESRS section code
    _ESRS_KEYWORD_MAP = [
        ('E1', ['environmental'], ['scope 1','scope 2','scope 3','co2','carbone','ges','ghg','émission','énergie','renouvelable','climatique','température']),
        ('E2', ['environmental'], ['pollution','polluant','rejet','atmosphérique','contamination','sol']),
        ('E3', ['environmental'], ['eau','hydrique','water','aqua','marine','ressource']),
        ('E4', ['environmental'], ['biodiversité','écosystème','site protégé','faune','flore']),
        ('E5', ['environmental'], ['déchet','recyclage','circulaire','réemploi','emballage']),
        ('S1', ['social'],        ['effectif','salarié','employé','etp','contrat','cdi','cdd','formation','accident','turnover','femme','parité','diversité','rémunération']),
        ('S2', ['social'],        ['fournisseur','chaîne','supply','sous-traitant','prestataire']),
        ('S3', ['social'],        ['communauté','local','territoire','riverain','mécénat']),
        ('S4', ['social'],        ['consommateur','client','satisfaction','réclamation']),
        ('G1', ['governance'],    []),  # all governance goes to G1
    ]

    def _map_to_esrs(self, pillar: str, metric_name: str, category: str) -> str:
        """Map a data entry to an ESRS section code using keyword matching."""
        text = (metric_name + ' ' + (category or '')).lower()
        for code, pillars, keywords in self._ESRS_KEYWORD_MAP:
            if pillar not in pillars:
                continue
            if not keywords:  # catch-all for the pillar
                return code
            if any(kw in text for kw in keywords):
                return code
        # Fallback by pillar
        fallback = {'environmental': 'E1', 'social': 'S1', 'governance': 'G1'}
        return fallback.get(pillar, 'E1')

    async def _collect_data(
        self,
        tenant_id: UUID,
        organization_id: Optional[UUID],
        year: Optional[int]
    ) -> Dict[str, Any]:
        # ── Fetch data entries (no year filter to include all available data) ──
        query = select(DataEntry).where(DataEntry.tenant_id == tenant_id)
        if organization_id:
            query = query.where(DataEntry.organization_id == organization_id)
        if year:
            # Include entries from target year AND previous year for completeness
            query = query.where(extract('year', DataEntry.period_start).in_([year, year - 1]))

        result = await self.db.execute(query)
        entries = result.scalars().all()

        # ── Fetch real ESG score from esg_scores table ──
        real_score = 0
        latest_period_months: Optional[int] = None
        real_rating = 'N/A'
        latest_e_score = 0.0
        latest_s_score = 0.0
        latest_g_score = 0.0
        score_history: list = []
        org_name = 'Organisation'
        try:
            async with self.db.begin_nested():
                score_q = select(ESGScore).where(ESGScore.tenant_id == tenant_id)
                if organization_id:
                    score_q = score_q.where(ESGScore.organization_id == organization_id)
                score_q = score_q.order_by(ESGScore.created_at.desc()).limit(1)
                score_res = await self.db.execute(score_q)
                esg_score = score_res.scalar_one_or_none()
                if esg_score:
                    real_score = round(esg_score.overall_score or 0)
                    real_rating = esg_score.rating or 'N/A'
                    latest_e_score = float(esg_score.environmental_score or 0)
                    latest_s_score = float(esg_score.social_score or 0)
                    latest_g_score = float(esg_score.governance_score or 0)
                    latest_period_months = getattr(esg_score, 'period_months', None)

                # Score history — last 12 entries for trend analysis
                hist_q = select(ESGScore).where(ESGScore.tenant_id == tenant_id)
                if organization_id:
                    hist_q = hist_q.where(ESGScore.organization_id == organization_id)
                hist_q = hist_q.order_by(ESGScore.calculation_date.desc()).limit(12)
                hist_res = await self.db.execute(hist_q)
                for s in reversed(hist_res.scalars().all()):  # chronological
                    score_history.append({
                        'date': s.calculation_date.isoformat() if s.calculation_date else None,
                        'overall': round(s.overall_score or 0, 1),
                        'environmental': round(s.environmental_score or 0, 1),
                        'social': round(s.social_score or 0, 1),
                        'governance': round(s.governance_score or 0, 1),
                        'rating': s.rating or '—',
                        'data_completeness': round(s.data_completeness or 0, 1),
                        'period_months': getattr(s, 'period_months', None),
                    })

                # Get org name: by ID if provided, else first org in tenant
                if organization_id:
                    org_q = select(Organization).where(Organization.id == organization_id)
                else:
                    org_q = select(Organization).where(Organization.tenant_id == tenant_id).limit(1)
                org_res = await self.db.execute(org_q)
                org = org_res.scalar_one_or_none()
                if org:
                    org_name = org.name
        except Exception as exc:
            logger.warning("collect: score fetch failed: %s", exc)

        # ── Materiality issues (Double materiality) ────────────────────────────
        # Each best-effort fetch runs in its own savepoint so a single failure
        # (table missing, schema drift, …) doesn't poison the outer transaction
        # and prevent the report_history INSERT downstream.
        materiality_issues: list = []
        try:
            async with self.db.begin_nested():
                from app.models.materiality import MaterialityIssue
                mat_q = select(MaterialityIssue).where(MaterialityIssue.tenant_id == tenant_id)
                mat_res = await self.db.execute(mat_q)
                for m in mat_res.scalars().all():
                    materiality_issues.append({
                        'name': m.name,
                        'category': m.category,
                        'financial_impact': float(m.financial_impact or 0),
                        'esg_impact': float(m.esg_impact or 0),
                        'is_material': bool(m.is_material),
                        'priority': m.priority or '—',
                        'stakeholders': m.stakeholders or '',
                        'description': m.description or '',
                    })
        except Exception as exc:
            logger.warning("collect: materiality fetch failed: %s", exc)

        # ── ESG Risks register ─────────────────────────────────────────────────
        esg_risks: list = []
        try:
            async with self.db.begin_nested():
                from app.models.materiality import ESGRisk
                risk_q = select(ESGRisk).where(ESGRisk.tenant_id == tenant_id).order_by(ESGRisk.risk_score.desc().nullslast())
                risk_res = await self.db.execute(risk_q)
                for r in risk_res.scalars().all():
                    esg_risks.append({
                        'title': r.title,
                        'category': r.category,
                        'probability': int(r.probability or 0),
                        'impact': int(r.impact or 0),
                        'risk_score': int(r.risk_score) if r.risk_score is not None else (int(r.probability or 0) * int(r.impact or 0)),
                        'severity': r.severity or '—',
                        'status': r.status or 'active',
                        'mitigation_status': r.mitigation_status or '—',
                        'responsible': r.responsible_person or '',
                        'mitigation_plan': r.mitigation_plan or '',
                        'target_date': r.target_date.isoformat() if r.target_date else None,
                    })
        except Exception as exc:
            logger.warning("collect: risks fetch failed: %s", exc)

        # ── Suppliers (Supply chain ESG) ───────────────────────────────────────
        suppliers: list = []
        try:
            async with self.db.begin_nested():
                from app.models.supplier import Supplier
                sup_q = select(Supplier).where(Supplier.tenant_id == tenant_id).order_by(Supplier.global_score.desc().nullslast()).limit(100)
                sup_res = await self.db.execute(sup_q)
                for s in sup_res.scalars().all():
                    suppliers.append({
                        'name': s.name,
                        'country': s.country or '—',
                        'category': s.category or '—',
                        'risk_level': s.risk_level or '—',
                        'status': s.status or '—',
                        'employees': int(s.employees) if s.employees else None,
                        'spend_k_eur': float(s.spend_k_eur) if s.spend_k_eur else None,
                        'global_score': round(float(s.global_score), 1) if s.global_score is not None else None,
                        'env_score': round(float(s.env_score), 1) if s.env_score is not None else None,
                        'social_score': round(float(s.social_score), 1) if s.social_score is not None else None,
                        'gov_score': round(float(s.gov_score), 1) if s.gov_score is not None else None,
                    })
        except Exception as exc:
            logger.warning("collect: suppliers fetch failed: %s", exc)

        # ── Indicator catalog (computed indicator_data) ────────────────────────
        indicator_catalog: list = []
        try:
            async with self.db.begin_nested():
                from app.models.indicator_data import IndicatorData
                from app.models.indicator import Indicator
                from sqlalchemy.orm import selectinload
                id_q = (
                    select(IndicatorData)
                    .where(IndicatorData.tenant_id == tenant_id)
                    .options(selectinload(IndicatorData.indicator))
                    .order_by(IndicatorData.date.desc())
                    .limit(500)
                )
                id_res = await self.db.execute(id_q)
                for d in id_res.scalars().all():
                    ind = d.indicator
                    indicator_catalog.append({
                        'date': d.date.isoformat() if d.date else None,
                        'code': getattr(ind, 'code', None) or '—',
                        'name': getattr(ind, 'name', None) or '—',
                        'value': float(d.value) if d.value is not None else None,
                        'unit': d.unit or getattr(ind, 'unit', '') or '',
                        'source': d.source or '—',
                        'is_verified': bool(d.is_verified),
                        'validation_status': d.validation_status or 'pending',
                    })
        except Exception as exc:
            logger.warning("collect: indicator catalog fetch failed: %s", exc)

        data_by_pillar: Dict[str, list] = {
            'environmental': [], 'social': [], 'governance': [],
        }
        by_section: Dict[str, list] = {s['code']: [] for s in ESRS_SECTIONS}

        for entry in entries:
            pillar = getattr(entry, 'pillar', None) or ''
            if pillar not in data_by_pillar:
                continue
            has_value = entry.value_numeric is not None or (entry.value_text and entry.value_text.strip())
            item = {
                'metric_name': entry.metric_name or '—',
                'value': entry.value_numeric if entry.value_numeric is not None else (entry.value_text or '—'),
                'unit': entry.unit or '',
                'category': entry.category or '',
                'verification_status': entry.verification_status or 'pending',
                'has_value': has_value,
            }
            data_by_pillar[pillar].append(item)
            # Map to ESRS section using keyword matching
            section_code = self._map_to_esrs(pillar, entry.metric_name or '', entry.category or '')
            if section_code in by_section:
                by_section[section_code].append(item)

        total = len(entries)
        with_data = sum(1 for e in entries if e.value_numeric is not None)
        pending = total - with_data

        # ── Calculated KPIs (automatic formulas engine) ────────────────────────
        # organization_id=None intentionnel : les données Calcul auto sont sauvegardées
        # sans org_id (niveau tenant), donc on cherche sans filtre organisation.
        calculated_kpis: dict = {}
        try:
            from app.services.calculation_service import CalculationService
            calc_service = CalculationService(self.db)
            calculated_kpis = await calc_service.calculate_metrics(
                tenant_id=tenant_id,
                organization_id=None,   # tenant-wide : données sans org_id
                year=year,
            )
        except Exception as exc:
            logger.warning("collect: calculated_kpis failed: %s", exc)

        return {
            'entries': data_by_pillar,
            'by_section': by_section,
            'org_name': org_name,
            'calculated_kpis': calculated_kpis,
            'stats': {
                'total_entries': total,
                'verified_count': with_data,
                'pending_count': pending,
                'score': real_score,
                'rating': real_rating,
                'environmental_score': round(latest_e_score, 1),
                'social_score': round(latest_s_score, 1),
                'governance_score': round(latest_g_score, 1),
                'period_months': latest_period_months,
                'by_pillar': {
                    'environmental': len(data_by_pillar['environmental']),
                    'social': len(data_by_pillar['social']),
                    'governance': len(data_by_pillar['governance']),
                },
            },
            'total_count': total,
            'score_history': score_history,
            'materiality_issues': materiality_issues,
            'esg_risks': esg_risks,
            'suppliers': suppliers,
            'indicator_catalog': indicator_catalog,
        }

    # ─── PDF Generation ───────────────────────────────────────────────────────

    async def _generate_pdf(self, report_type: str, data: Dict[str, Any], year: int) -> bytes:
        buffer = io.BytesIO()
        report_name = self.REPORT_TYPES[report_type]['name']
        stats = data['stats']

        # Dispatch CSRD to dedicated generator
        if report_type == 'csrd':
            return await self._generate_csrd_pdf(data, year)

        # Dispatch DPEF to dedicated generator
        if report_type == 'dpef':
            return await self._generate_dpef_pdf(data, year)

        story = []

        # Cover page
        story += self._make_cover(report_name, year, stats, data.get('org_name', 'Organisation'), data=data)
        story.append(PageBreak())

        # Executive summary
        story += self._make_executive_summary(stats, data)
        story.append(PageBreak())

        # KPIs Calculés Automatiquement (ESRS)
        story += self._make_calculated_kpis_section(data.get('calculated_kpis', {}))
        story.append(PageBreak())

        # ESRS completeness
        story += self._make_esrs_completeness(data['by_section'])
        story.append(PageBreak())

        # Carbon-specific section (scopes donut + breakdown) — for carbon & TCFD reports
        if report_type in ('carbon', 'tcfd'):
            carbon_section = self._make_carbon_section(data['entries'])
            if carbon_section:
                story += carbon_section
                story.append(PageBreak())

        # Pillar data tables
        story += self._make_pillar_data(data['entries'])

        # Annexes (matérialité, risques, fournisseurs, historique scores)
        story.append(PageBreak())
        story += self._make_annexes(data)

        # Compliance page
        story += self._make_compliance_page(year)

        first_page_cb, later_pages_cb = _make_header_footer(report_name, year)

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=22*mm,
            bottomMargin=22*mm,
        )

        doc.build(story, onFirstPage=first_page_cb, onLaterPages=later_pages_cb)

        buffer.seek(0)
        return buffer.read()

    # ─── CSRD Dedicated Generator ─────────────────────────────────────────────

    async def _generate_csrd_pdf(self, data: Dict[str, Any], year: int) -> bytes:
        """Generate a structured CSRD compliance report with ESRS sections."""
        buffer = io.BytesIO()
        stats = data['stats']
        org_name = data.get('org_name', 'Organisation')
        by_section = data.get('by_section', {})

        story = []

        # 1. Cover page (reuse existing)
        story += self._make_cover('Rapport CSRD', year, stats, org_name, data=data)
        story.append(PageBreak())

        # 2. Table of contents
        story += self._make_csrd_toc()
        story.append(PageBreak())

        # 3. Executive Summary
        story += self._make_executive_summary(stats, data)
        story.append(PageBreak())

        # 3b. KPIs Calculés Automatiquement (ESRS)
        story += self._make_calculated_kpis_section(data.get('calculated_kpis', {}))
        story.append(PageBreak())

        # 4. ESRS sections one by one (E, S, G)
        story += self._make_csrd_esrs_sections(by_section, data['entries'])
        story.append(PageBreak())

        # 5. Gap analysis table
        story += self._make_csrd_gap_table(by_section)
        story.append(PageBreak())

        # 6. Double Materiality Analysis (DMA) — ESRS 2 IRO-1/IRO-2
        dma_section = self._make_dma_page(data)
        if dma_section:
            story += dma_section
            story.append(PageBreak())

        # 7. Annexes (matérialité, risques, fournisseurs, historique scores)
        story += self._make_annexes(data)

        # 8. Compliance page
        story += self._make_compliance_page(year)

        first_page_cb, later_pages_cb = _make_header_footer('Rapport CSRD', year)

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=22*mm,
            bottomMargin=22*mm,
        )
        doc.build(story, onFirstPage=first_page_cb, onLaterPages=later_pages_cb)
        buffer.seek(0)
        return buffer.read()

    # ─── DPEF Dedicated Generator ─────────────────────────────────────────────

    # DPEF Art. L.225-102-1 — three pillars with predefined disclosure topics
    DPEF_SECTIONS = {
        'social': {
            'label': 'Informations sociales',
            'color': C_BLUE,
            'bg': C_BLUE2,
            'topics': [
                'Effectif total et répartition', 'Embauches et licenciements',
                'Rémunérations et évolution', 'Organisation du temps de travail',
                'Absentéisme', 'Relations sociales',
                'Conditions de santé et sécurité', 'Accidents du travail et maladies professionnelles',
                'Formation professionnelle', 'Égalité de traitement',
                'Promotion et respect des conventions OIT', 'Handicap',
            ],
        },
        'environmental': {
            'label': 'Informations environnementales',
            'color': C_GREEN,
            'bg': C_GREEN2,
            'topics': [
                'Politique générale environnementale', 'Pollution et gestion des déchets',
                'Utilisation durable des ressources', 'Changement climatique',
                'Protection de la biodiversité', 'Émissions de GES (Scopes 1/2/3)',
                'Consommation d\'énergie', 'Consommation d\'eau',
                'Utilisation des sols', 'Mesures de prévention et réparation',
            ],
        },
        'societal': {
            'label': 'Informations sociétales',
            'color': C_PURPLE,
            'bg': C_PURPLE2,
            'topics': [
                'Engagements sociétaux en faveur du développement durable',
                'Sous-traitance et fournisseurs', 'Loyauté des pratiques',
                'Actions en faveur des droits de l\'homme',
                'Lutte contre la corruption', 'Lutte contre l\'évasion fiscale',
                'Actions en faveur de l\'économie circulaire',
                'Actions en faveur de la lutte contre le gaspillage alimentaire',
            ],
        },
    }

    async def _generate_dpef_pdf(self, data: Dict[str, Any], year: int) -> bytes:
        """Generate a structured DPEF (Art. L.225-102-1) report.

        Mirrors the frontend DPEFReport.tsx but uses ReportLab for professional
        quality: cover page → TOC → social → environmental → societal → carbon
        section → compliance page.
        """
        buffer = io.BytesIO()
        stats = data['stats']
        org_name = data.get('org_name', 'Organisation')
        entries = data.get('entries', {})
        story = []

        # 1. Cover
        story += self._make_cover('DPEF — Déclaration de performance extra-financière', year, stats, org_name, data=data)
        story.append(PageBreak())

        # 2. TOC
        story += self._section_header('Sommaire', C_NAVY)
        dpef_toc = [
            ('1.', 'Résumé exécutif',             '3'),
            ('2.', 'Informations sociales',         '4'),
            ('3.', 'Informations environnementales', '6'),
            ('4.', 'Informations sociétales',       '8'),
            ('5.', 'Bilan carbone (Scopes 1/2/3)',  '9'),
            ('6.', 'Conformité & méthodologie',     '10'),
        ]
        toc_rows = []
        for num, title, page in dpef_toc:
            toc_rows.append([
                Paragraph(f'<b>{num}</b>', ParagraphStyle('dtn', fontSize=10, textColor=C_GREEN)),
                Paragraph(title, ParagraphStyle('dtt', fontSize=10, textColor=C_NAVY)),
                Paragraph(page, ParagraphStyle('dtp', fontSize=10, textColor=C_GRAY, alignment=2)),
            ])
        t = Table(toc_rows, colWidths=[10*mm, 140*mm, 20*mm])
        t.setStyle(TableStyle([
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [C_WHITE, C_LGRAY]),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (0, -1), 8),
            ('LINEBELOW', (0, -1), (-1, -1), 1.5, C_GREEN),
        ]))
        story.append(t)
        story.append(PageBreak())

        # 3. Executive summary
        story += self._make_executive_summary(stats, data)
        story.append(PageBreak())

        # 3b. KPIs Calculés Automatiquement (ESRS)
        story += self._make_calculated_kpis_section(data.get('calculated_kpis', {}))
        story.append(PageBreak())

        # 4. Three DPEF pillars
        pillar_map = {'social': 'social', 'environmental': 'environmental', 'societal': 'governance'}
        for pillar_key, section_cfg in self.DPEF_SECTIONS.items():
            story += self._section_header(section_cfg['label'], section_cfg['color'], size=14)

            story.append(Paragraph(
                f"<i>Art. L.225-102-1 du Code de commerce — {section_cfg['label']}</i>",
                ParagraphStyle('dpef_ref', fontSize=8, textColor=C_GRAY, spaceAfter=6, leading=11)
            ))

            # Map topics against available data entries
            data_key = pillar_map.get(pillar_key, pillar_key)
            pillar_entries = entries.get(data_key, [])

            rows = [['Thématique DPEF', 'Données disponibles', 'Statut']]
            for topic in section_cfg['topics']:
                # Try to match entries to this topic by keyword overlap
                topic_lower = topic.lower()
                matched = [
                    e for e in pillar_entries
                    if any(kw in (str(e.get('metric_name') or '') + ' ' + str(e.get('category') or '')).lower()
                           for kw in topic_lower.split()[:3])
                ]
                if matched:
                    val_str = ', '.join(
                        f"{e.get('metric_name', '—')}: {e.get('value', '—')} {e.get('unit', '')}"
                        for e in matched[:2]
                    )
                    if len(val_str) > 60:
                        val_str = val_str[:57] + '…'
                    status = 'Renseigné'
                else:
                    val_str = '—'
                    status = 'À compléter'
                rows.append([topic[:40], val_str, status])

            dpef_table = Table(rows, colWidths=[52*mm, 72*mm, 24*mm])
            dpef_table.setStyle(TableStyle([
                ('BACKGROUND',     (0, 0), (-1, 0), section_cfg['color']),
                ('TEXTCOLOR',      (0, 0), (-1, 0), colors.white),
                ('FONTNAME',       (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',       (0, 0), (-1, 0), 8),
                ('FONTSIZE',       (0, 1), (-1, -1), 7.5),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [C_WHITE, C_LGRAY]),
                ('GRID',           (0, 0), (-1, -1), 0.25, C_MGRAY),
                ('ALIGN',          (2, 0), (2, -1), 'CENTER'),
                ('VALIGN',         (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING',     (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING',  (0, 0), (-1, -1), 3),
                ('LEFTPADDING',    (0, 0), (0, -1), 6),
                ('LEFTPADDING',    (1, 0), (1, -1), 6),
            ]))
            story.append(dpef_table)
            story.append(Spacer(1, 4*mm))

            # Summary stats for this pillar
            filled = sum(1 for r in rows[1:] if r[2] == 'Renseigné')
            total_topics = len(section_cfg['topics'])
            pct = round(filled / total_topics * 100) if total_topics else 0
            color_bar = section_cfg['color']
            story.append(Paragraph(
                f'<b>{filled}/{total_topics}</b> thématiques renseignées ({pct} %)',
                ParagraphStyle('dpef_pct', fontSize=9, textColor=C_NAVY, spaceAfter=2)
            ))
            story.append(PageBreak())

        # 5. Carbon section (reuse existing)
        carbon = self._make_carbon_section(entries)
        if carbon:
            story += carbon
            story.append(PageBreak())

        # 6. Compliance page
        story += self._make_compliance_page(year)

        first_page_cb, later_pages_cb = _make_header_footer('DPEF', year)
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=22*mm,
            bottomMargin=22*mm,
        )
        doc.build(story, onFirstPage=first_page_cb, onLaterPages=later_pages_cb)
        buffer.seek(0)
        return buffer.read()

    def _make_csrd_toc(self) -> list:
        """Table of contents page."""
        story = []
        story += self._section_header('Sommaire', C_NAVY)

        toc_items = [
            ('1.', 'Résumé exécutif',                          '3'),
            ('2.', 'Pilier Environnemental (ESRS E1–E5)',       '4'),
            ('3.', 'Pilier Social (ESRS S1–S4)',                '6'),
            ('4.', 'Pilier Gouvernance (ESRS G1)',              '8'),
            ('5.', 'Analyse des écarts (Gap Analysis)',         '9'),
            ('6.', 'Double matérialité (DMA — ESRS 2 IRO)',    '10'),
            ('7.', 'Cadres de référence & Conformité',          '11'),
        ]

        rows = []
        for num, title, page in toc_items:
            rows.append([
                Paragraph(f'<b>{num}</b>', ParagraphStyle('tn', fontSize=10, textColor=C_GREEN)),
                Paragraph(title, ParagraphStyle('tt', fontSize=10, textColor=C_NAVY)),
                Paragraph(page, ParagraphStyle('tp', fontSize=10, textColor=C_GRAY, alignment=2)),
            ])

        t = Table(rows, colWidths=[10*mm, 140*mm, 20*mm])
        t.setStyle(TableStyle([
            ('ROWBACKGROUNDS', (0,0), (-1,-1), [C_WHITE, C_LGRAY]),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (0,-1), 8),
            ('LINEBELOW', (0,-1), (-1,-1), 1.5, C_GREEN),
        ]))
        story.append(t)
        return story

    def _make_csrd_esrs_sections(self, by_section: dict, entries: dict) -> list:
        """Generate one section per ESRS standard with data table."""
        story = []

        pillar_groups = [
            ('environmental', 'Pilier Environnemental', C_GREEN, C_GREEN2, [s for s in ESRS_SECTIONS if s['pillar'] == 'environmental']),
            ('social',        'Pilier Social',          C_BLUE,  C_BLUE2,  [s for s in ESRS_SECTIONS if s['pillar'] == 'social']),
            ('governance',    'Pilier Gouvernance',     C_PURPLE, C_PURPLE2, [s for s in ESRS_SECTIONS if s['pillar'] == 'governance']),
        ]

        for pillar_key, pillar_label, p_color, p_bg, sections in pillar_groups:
            story += self._section_header(pillar_label, p_color, size=14)

            pillar_entries = entries.get(pillar_key, [])
            pillar_total = len(pillar_entries)

            story.append(Paragraph(
                f'{pillar_total} indicateur(s) collecté(s) pour ce pilier.',
                ParagraphStyle('intro', fontSize=9, textColor=C_GRAY, spaceAfter=6)
            ))

            for sec in sections:
                sec_entries = by_section.get(sec['code'], [])
                count = len(sec_entries)
                pct = min(round(count / 5 * 100), 100) if count else 0

                # Section sub-header
                sec_header = Table(
                    [[
                        Paragraph(f'<font color="white"><b>{sec["code"]}</b></font>',
                                   ParagraphStyle('sh', fontSize=9, alignment=1)),
                        Paragraph(f'<font color="white"><b>{sec["label"]}</b></font>',
                                   ParagraphStyle('sl', fontSize=9)),
                        Paragraph(f'<font color="white">{pct}% couvert</font>',
                                   ParagraphStyle('sp', fontSize=8, alignment=2)),
                    ]],
                    colWidths=[16*mm, 120*mm, 34*mm],
                )
                sec_header.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), p_color),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ('TOPPADDING', (0,0), (-1,-1), 5),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                    ('LEFTPADDING', (0,0), (-1,-1), 8),
                    ('RIGHTPADDING', (0,0), (-1,-1), 8),
                ]))
                story.append(sec_header)

                if sec_entries:
                    # Data table for this section
                    rows = [['Métrique', 'Valeur', 'Unité', 'Statut']]
                    for entry in sec_entries[:15]:
                        status_text = '✓' if entry.get('verification_status') == 'verified' else '⏳'
                        rows.append([
                            str(entry.get('metric_name', ''))[:55],
                            str(entry.get('value', '—'))[:20],
                            str(entry.get('unit', ''))[:12],
                            status_text,
                        ])

                    dt = Table(rows, colWidths=[88*mm, 32*mm, 24*mm, 26*mm])
                    dt.setStyle(TableStyle([
                        ('BACKGROUND',    (0,0), (-1,0), p_bg),
                        ('TEXTCOLOR',     (0,0), (-1,0), p_color),
                        ('FONTNAME',      (0,0), (-1,0), 'Helvetica-Bold'),
                        ('FONTSIZE',      (0,0), (-1,0), 8),
                        ('FONTSIZE',      (0,1), (-1,-1), 7.5),
                        ('ROWBACKGROUNDS',(0,1), (-1,-1), [C_WHITE, C_LGRAY]),
                        ('GRID',          (0,0), (-1,-1), 0.2, C_MGRAY),
                        ('ALIGN',         (1,0), (-1,-1), 'CENTER'),
                        ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
                        ('TOPPADDING',    (0,0), (-1,-1), 4),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                        ('LEFTPADDING',   (0,0), (0,-1), 8),
                    ]))
                    story.append(dt)
                else:
                    # No data message
                    no_data = Table(
                        [[Paragraph('⚠ Aucune donnée collectée pour cette section.',
                                     ParagraphStyle('nd', fontSize=8, textColor=C_AMBER))]],
                        colWidths=[170*mm],
                    )
                    no_data.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,-1), C_AMBER2),
                        ('TOPPADDING', (0,0), (-1,-1), 6),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                        ('LEFTPADDING', (0,0), (-1,-1), 10),
                    ]))
                    story.append(no_data)

                story.append(Spacer(1, 3*mm))

            story.append(PageBreak())

        return story

    def _make_csrd_gap_table(self, by_section: dict) -> list:
        """Gap analysis table: ESRS standard | Coverage | Status | Recommended action."""
        story = []
        story += self._section_header('Analyse des écarts ESRS (Gap Analysis)', C_NAVY)
        story.append(Paragraph(
            'Récapitulatif de la couverture des exigences ESRS et actions recommandées.',
            ParagraphStyle('sub', fontSize=9, textColor=C_GRAY, spaceAfter=8)
        ))

        rows = [['Standard', 'Section', 'Données', 'Couverture', 'Statut', 'Action prioritaire']]

        for sec in ESRS_SECTIONS:
            count = len(by_section.get(sec['code'], []))
            pct = min(round(count / 5 * 100), 100) if count else 0

            if pct >= 80:
                status = '✓ Prêt'
                status_color = C_GREEN
                action = 'Maintenir & vérifier'
            elif pct >= 30:
                status = '⚠ Partiel'
                status_color = C_AMBER
                action = 'Compléter les données manquantes'
            else:
                status = '✗ Manquant'
                status_color = C_RED
                action = 'Collecte de données requise'

            # Color code the pillar
            if sec['pillar'] == 'environmental':
                code_color = C_GREEN
            elif sec['pillar'] == 'social':
                code_color = C_BLUE
            else:
                code_color = C_PURPLE

            rows.append([
                Paragraph(f'<font color="{_hex(code_color)}"><b>{sec["code"]}</b></font>',
                           ParagraphStyle('gc', fontSize=8)),
                Paragraph(sec['label'][:35],
                           ParagraphStyle('gl', fontSize=8, textColor=C_NAVY)),
                Paragraph(str(count),
                           ParagraphStyle('gd', fontSize=8, alignment=1)),
                self._progress_bar(pct, code_color, width=28*mm, height=4*mm),
                Paragraph(f'<font color="{_hex(status_color)}"><b>{status}</b></font>',
                           ParagraphStyle('gs', fontSize=7, alignment=1)),
                Paragraph(action,
                           ParagraphStyle('ga', fontSize=7, textColor=C_GRAY)),
            ])

        t = Table(rows, colWidths=[14*mm, 46*mm, 14*mm, 30*mm, 18*mm, 48*mm])
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0,0), (-1,0), C_NAVY),
            ('TEXTCOLOR',     (0,0), (-1,0), C_WHITE),
            ('FONTNAME',      (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE',      (0,0), (-1,0), 8),
            ('ROWBACKGROUNDS',(0,1), (-1,-1), [C_WHITE, C_LGRAY]),
            ('GRID',          (0,0), (-1,-1), 0.3, C_MGRAY),
            ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING',    (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('LEFTPADDING',   (0,0), (-1,-1), 6),
            ('ALIGN',         (2,0), (2,-1), 'CENTER'),
            ('ALIGN',         (4,0), (4,-1), 'CENTER'),
        ]))
        story.append(t)

        # Summary box
        story.append(Spacer(1, 6*mm))
        ready = sum(1 for s in ESRS_SECTIONS if min(round(len(by_section.get(s['code'], [])) / 5 * 100), 100) >= 80)
        partial = sum(1 for s in ESRS_SECTIONS if 30 <= min(round(len(by_section.get(s['code'], [])) / 5 * 100), 100) < 80)
        missing = len(ESRS_SECTIONS) - ready - partial

        summary = Table(
            [[
                self._kpi_cell(str(ready), 'Standards prêts', C_GREEN, C_GREEN2),
                self._kpi_cell(str(partial), 'Standards partiels', C_AMBER, C_AMBER2),
                self._kpi_cell(str(missing), 'Standards manquants', C_RED, colors.HexColor('#fee2e2')),
                self._kpi_cell(f'{round(ready/len(ESRS_SECTIONS)*100)}%', 'Conformité globale', C_BLUE, C_BLUE2),
            ]],
            colWidths=[40*mm, 40*mm, 40*mm, 40*mm],
        )
        summary.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(summary)

        return story

    # ─── Cover Page ───────────────────────────────────────────────────────────

    def _make_cover(self, report_name: str, year: int, stats: dict, org_name: str = 'Organisation', data: Optional[Dict[str, Any]] = None) -> list:
        """
        Cover page — modern, overlap-free design.

        Layout (top → bottom):
          ┌─────────────────────────────────────────────────┐
          │  Dark navy header  (logo row + subtitle row)    │
          ├─────────────────────────────────────────────────┤
          │  5 mm green accent bar                          │
          ├─────────────────────────────────────────────────┤
          │  24 mm breathing space                          │
          │  Org name (green, 13 pt, bold)                  │
          │  6 mm space                                     │
          │  Report type badge pill                         │
          │  6 mm space                                     │
          │  Report title (28 pt navy bold)   leading=36   │
          │  6 mm space                                     │
          │  "Exercice YYYY" (14 pt gray)                   │
          │  8 mm space                                     │
          │  Green HR line (50 % width)                     │
          │  6 mm space                                     │
          │  Subtitle caption (9 pt gray)                   │
          │  14 mm space                                    │
          │  KPI strip (4 cards, 2 rows each)               │
          │  12 mm space                                    │
          │  Generation info (8 pt gray)                    │
          └─────────────────────────────────────────────────┘

        Key fixes vs. previous version:
          • "ESG Flow" (34 pt) and subtitle split into TWO separate Table rows
            so leading never fights font size → zero overlap.
          • spaceAfter values increased to proper optical spacing.
          • _kpi_cell uses two separate rows (value / label) instead of
            a single Paragraph with <br/> and an insufficient leading.
        """
        story = []

        # ── 1. Navy header: TWO rows so leading is irrelevant ─────────────────
        header_logo = Paragraph(
            'ESG Flow',
            ParagraphStyle('hdr_logo', fontSize=34, textColor=C_WHITE,
                           alignment=1, fontName='Helvetica-Bold', leading=42)
        )
        header_sub = Paragraph(
            'Plateforme de reporting ESG',
            ParagraphStyle('hdr_sub', fontSize=11,
                           textColor=colors.HexColor('#a7f3d0'),
                           alignment=1, leading=15)
        )
        cover_header = Table(
            [[header_logo], [header_sub]],
            colWidths=[170*mm],
        )
        cover_header.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), C_NAVY),
            ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            # logo row: top=26 bottom=4
            ('TOPPADDING',    (0, 0), (0,  0),  26),
            ('BOTTOMPADDING', (0, 0), (0,  0),   4),
            # subtitle row: top=0 bottom=26
            ('TOPPADDING',    (0, 1), (0,  1),   0),
            ('BOTTOMPADDING', (0, 1), (0,  1),  26),
        ]))
        story.append(cover_header)

        # ── 2. Green accent bar (5 mm) ────────────────────────────────────────
        accent = Table([['']], colWidths=[170*mm], rowHeights=[5])
        accent.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), C_GREEN)]))
        story.append(accent)

        # ── 3. Breathing space ────────────────────────────────────────────────
        story.append(Spacer(1, 20*mm))

        # ── 4. Organisation name ──────────────────────────────────────────────
        story.append(Paragraph(
            org_name.upper(),
            ParagraphStyle('cov_org', fontSize=13, textColor=C_GREEN,
                           alignment=1, fontName='Helvetica-Bold', leading=18,
                           spaceAfter=6)
        ))

        # ── 5. Report type pill ───────────────────────────────────────────────
        pill = Table(
            [[Paragraph(
                report_name,
                ParagraphStyle('pill_txt', fontSize=9, textColor=C_WHITE,
                               alignment=1, leading=12)
            )]],
            colWidths=[60*mm],
        )
        pill.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), C_GREEN),
            ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
            ('TOPPADDING',    (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('ROUNDEDCORNERS', [6]),
        ]))
        # Centre the pill using a 3-column wrapper (left pad | pill | right pad)
        pill_row = Table(
            [[Paragraph('', ParagraphStyle('x')), pill, Paragraph('', ParagraphStyle('x'))]],
            colWidths=[55*mm, 60*mm, 55*mm],
        )
        pill_row.setStyle(TableStyle([
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(pill_row)

        # ── 6. Report title (separate paragraph — no overlap risk) ────────────
        story.append(Spacer(1, 8*mm))
        story.append(Paragraph(
            report_name.upper(),
            ParagraphStyle('cov_title', fontSize=28, textColor=C_NAVY,
                           alignment=1, fontName='Helvetica-Bold',
                           leading=36, spaceAfter=0)
        ))

        # ── 7. Year (separate paragraph, enough top space) ───────────────────
        story.append(Spacer(1, 6*mm))
        story.append(Paragraph(
            f'Exercice {year}',
            ParagraphStyle('cov_year', fontSize=14, textColor=C_GRAY,
                           alignment=1, leading=20, spaceAfter=0)
        ))

        # ── 8. Green divider ──────────────────────────────────────────────────
        story.append(Spacer(1, 8*mm))
        story.append(HRFlowable(
            width='45%', thickness=2, color=C_GREEN,
            lineCap='round', spaceAfter=0
        ))

        # ── 9. Subtitle caption ───────────────────────────────────────────────
        story.append(Spacer(1, 5*mm))
        story.append(Paragraph(
            'Rapport de durabilité conforme aux exigences CSRD / ESRS',
            ParagraphStyle('cov_sub', fontSize=9, textColor=C_GRAY,
                           alignment=1, leading=13)
        ))

        # ── 10. HERO score gauge + sparkline + 3 supporting KPIs ──────────────
        story.append(Spacer(1, 10*mm))
        score  = stats.get('score', 0)
        rating = stats.get('rating', 'N/A')
        history = ((data or {}).get('score_history')
                   or stats.get('score_history') or [])
        gauge = self._make_score_gauge(score, rating, width=60*mm, height=60*mm)
        spark = self._make_score_sparkline(history, width=88*mm, height=24*mm) if history else Spacer(1, 0)

        # Right column: 3 stacked supporting KPIs
        right_col = Table(
            [[self._kpi_cell(str(stats.get('total_entries', 0)), 'Indicateurs',  C_NAVY,  C_LGRAY)],
             [self._kpi_cell(str(stats.get('verified_count', 0)), 'Vérifiés',     C_GREEN, C_GREEN2)],
             [self._kpi_cell(str(stats.get('coverage_pct',
                              int(stats.get('verified_count', 0) /
                                  max(1, stats.get('total_entries', 1)) * 100))) + '%',
                              'Couverture', C_BLUE, C_BLUE2)]],
            colWidths=[42*mm],
        )
        right_col.setStyle(TableStyle([
            ('TOPPADDING',    (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))

        # Compose hero block: gauge on left, sparkline+KPIs on right
        hero = Table(
            [[gauge, Table([[spark], [right_col]],
                           colWidths=[88*mm])]],
            colWidths=[70*mm, 95*mm],
        )
        hero.setStyle(TableStyle([
            ('VALIGN',     (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN',      (0, 0), (0,  0),  'CENTER'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(hero)

        # ── 11. Window label (data window the score was computed on) ──────────
        period_months = stats.get('period_months')
        if period_months:
            story.append(Spacer(1, 6*mm))
            story.append(Paragraph(
                f"<b>Fenêtre de données :</b> {_humanize_period(period_months)}",
                ParagraphStyle('cov_win', fontSize=9, textColor=C_NAVY,
                               alignment=1, leading=12)
            ))

        # ── 12. Generation info ───────────────────────────────────────────────
        story.append(Spacer(1, 8*mm))
        story.append(Paragraph(
            f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} · ESG Flow v2.0.0",
            ParagraphStyle('cov_date', fontSize=8, textColor=C_GRAY,
                           alignment=1, leading=11)
        ))

        return story

    def _kpi_cell(self, value: str, label: str, text_color, bg_color) -> Table:
        """
        KPI card with value and label in SEPARATE rows — no leading conflict.
        Previous version used a single Paragraph with <br/> and leading=16,
        which caused overlap when value font was size 20.
        """
        t = Table(
            [
                [Paragraph(
                    f'<b>{value}</b>',
                    ParagraphStyle('kpi_val', fontSize=20, textColor=text_color,
                                   alignment=1, leading=26)
                )],
                [Paragraph(
                    label,
                    ParagraphStyle('kpi_lbl', fontSize=8, textColor=C_GRAY,
                                   alignment=1, leading=11)
                )],
            ],
            colWidths=[38*mm],
            rowHeights=[14*mm, 8*mm],
        )
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), bg_color),
            ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN',        (0, 0), (0,  0),  'BOTTOM'),   # value sticks to bottom
            ('VALIGN',        (0, 1), (0,  1),  'TOP'),      # label sticks to top
            ('ROUNDEDCORNERS', [6]),
            ('TOPPADDING',    (0, 0), (0,  0),   8),
            ('BOTTOMPADDING', (0, 0), (0,  0),   2),
            ('TOPPADDING',    (0, 1), (0,  1),   2),
            ('BOTTOMPADDING', (0, 1), (0,  1),   8),
            ('LEFTPADDING',   (0, 0), (-1, -1),  4),
            ('RIGHTPADDING',  (0, 0), (-1, -1),  4),
        ]))
        return t

    # ─── Executive Summary ────────────────────────────────────────────────────

    def _make_executive_summary(self, stats: dict, data: Optional[Dict[str, Any]] = None) -> list:
        story = []
        story += self._section_header('Résumé Exécutif', C_NAVY)

        # ── Window banner — make the data window unambiguous ───────────────
        period_months = stats.get('period_months')
        if period_months:
            from datetime import datetime as _dt, timedelta as _td, date as _d
            window_end   = _d.today()
            window_start = window_end - _td(days=int(period_months) * 30)
            story.append(self._make_insight_card(
                title='Fenêtre de calcul du score',
                message=(
                    f"Ce score a été calculé sur les données disponibles "
                    f"<b>du {window_start.strftime('%d/%m/%Y')} au {window_end.strftime('%d/%m/%Y')}</b> "
                    f"({_humanize_period(period_months)}). "
                    f"Les valeurs antérieures à cette fenêtre ne sont pas prises en compte."
                ),
                color=C_NAVY, bg_color=C_LGRAY, icon='⏱',
            ))
            story.append(Spacer(1, 5*mm))

        # Insight cards (data-driven contextual narrative)
        if data:
            for card in self._build_insights(data):
                story.append(card)
                story.append(Spacer(1, 3*mm))
            if self._build_insights(data):
                story.append(Spacer(1, 4*mm))

        score = stats['score']

        # Two-column layout: gauge left, pillar bars right
        gauge = self._make_gauge(score)
        pillar_table = self._make_pillar_bars(stats['by_pillar'])

        layout = Table(
            [[gauge, pillar_table]],
            colWidths=[75*mm, 95*mm],
        )
        layout.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(layout)
        story.append(Spacer(1, 8*mm))

        # ─── 3 donuts E/S/G — pillar performance at a glance ─────────────
        # Pulls latest pillar scores from history; falls back to global score
        # for any pillar without data.
        if data:
            history = data.get('score_history') or []
            latest = history[-1] if history else {}
            env_s = float(latest.get('environmental') or score)
            soc_s = float(latest.get('social')        or score)
            gov_s = float(latest.get('governance')    or score)
            story += self._section_header("Performance par pilier ESG", C_NAVY, size=12)
            story.append(self._make_pillar_donuts(env_s, soc_s, gov_s,
                                                  width=170*mm, height=58*mm))
            story.append(Spacer(1, 6*mm))

        # ─── ESRS completeness bars ──────────────────────────────────────
        if data and data.get('by_section'):
            story += self._section_header("Complétude des standards ESRS", C_NAVY, size=12)
            story.append(self._make_completeness_bars(data['by_section'],
                                                     width=170*mm, height=62*mm))
            story.append(Spacer(1, 6*mm))

        # ─── Score evolution chart (if enough history) ──────────────────
        if data:
            history = data.get('score_history') or []
            if len(history) >= 2:
                story += self._section_header("Évolution du score ESG (12 derniers points)", C_NAVY, size=12)
                story.append(self._make_score_line_chart(history))
                story.append(Spacer(1, 6*mm))

        # Stats summary table
        story += self._section_header('Données collectées', C_GREEN, size=13)
        rows = [
            ['Indicateur', 'Valeur', 'Statut'],
            ['Total indicateurs ESG',    str(stats['total_entries']),  ''],
            ['Données vérifiées',         str(stats['verified_count']), '✓'],
            ['Données en attente',        str(stats['pending_count']),  '⏳'],
            ['Taux de vérification',      f"{score}%",                  ''],
            ['Pilier Environnemental',    str(stats['by_pillar']['environmental']), ''],
            ['Pilier Social',             str(stats['by_pillar']['social']),        ''],
            ['Pilier Gouvernance',        str(stats['by_pillar']['governance']),    ''],
        ]
        t = Table(rows, colWidths=[90*mm, 40*mm, 40*mm])
        t.setStyle(TableStyle([
            ('BACKGROUND',   (0,0), (-1,0), C_NAVY),
            ('TEXTCOLOR',    (0,0), (-1,0), C_WHITE),
            ('FONTNAME',     (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE',     (0,0), (-1,0), 10),
            ('FONTSIZE',     (0,1), (-1,-1), 9),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [C_WHITE, C_LGRAY]),
            ('GRID',         (0,0), (-1,-1), 0.3, C_MGRAY),
            ('ALIGN',        (1,0), (-1,-1), 'CENTER'),
            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING',   (0,0), (-1,-1), 6),
            ('BOTTOMPADDING',(0,0), (-1,-1), 6),
            ('LEFTPADDING',  (0,0), (0,-1), 10),
        ]))
        story.append(t)
        return story

    def _make_gauge(self, score: int) -> Drawing:
        d = Drawing(70*mm, 80*mm)

        cx, cy, r = 35*mm, 45*mm, 28*mm

        # Background circle (gray)
        pie = Pie()
        pie.x = cx - r
        pie.y = cy - r
        pie.width  = 2 * r
        pie.height = 2 * r
        pie.data = [score, 100 - score]
        pie.startAngle = 90
        pie.direction = 'clockwise'
        pie.slices[0].fillColor = C_GREEN if score >= 70 else (C_AMBER if score >= 40 else C_RED)
        pie.slices[0].strokeColor = colors.white
        pie.slices[0].strokeWidth = 2
        pie.slices[1].fillColor = C_MGRAY
        pie.slices[1].strokeColor = colors.white
        pie.slices[1].strokeWidth = 2
        pie.innerRadiusFraction = 0.65
        d.add(pie)

        # Center text
        d.add(String(cx, cy + 5, f'{score}%',
                     fontSize=18, fontName='Helvetica-Bold',
                     fillColor=C_NAVY, textAnchor='middle'))
        d.add(String(cx, cy - 10, 'Score global',
                     fontSize=7, fontName='Helvetica',
                     fillColor=C_GRAY, textAnchor='middle'))

        # Label below
        d.add(String(cx, 8*mm, 'Complétude ESRS',
                     fontSize=8, fontName='Helvetica',
                     fillColor=C_GRAY, textAnchor='middle'))
        return d

    def _make_pillar_bars(self, by_pillar: dict) -> Table:
        env_count = by_pillar['environmental']
        soc_count = by_pillar['social']
        gov_count = by_pillar['governance']
        total = max(env_count + soc_count + gov_count, 1)

        rows = [
            [Paragraph('<b>Répartition par pilier</b>',
                       ParagraphStyle('ph', fontSize=10, textColor=C_NAVY))],
        ]

        for label, count, color in [
            ('Environnemental', env_count, C_GREEN),
            ('Social',          soc_count, C_BLUE),
            ('Gouvernance',     gov_count, C_PURPLE),
        ]:
            pct = round(count / total * 100)
            bar = self._progress_bar(pct, color, width=85*mm)
            sub = Table(
                [[Paragraph(label, ParagraphStyle('bl', fontSize=9, textColor=C_NAVY)),
                  Paragraph(f'<font color="{_hex(color)}"><b>{count}</b></font>',
                            ParagraphStyle('bv', fontSize=9, alignment=2))]],
                colWidths=[65*mm, 20*mm],
            )
            sub.setStyle(TableStyle([
                ('TOPPADDING', (0,0),(-1,-1), 2),
                ('BOTTOMPADDING', (0,0),(-1,-1), 2),
                ('LEFTPADDING', (0,0),(-1,-1), 0),
                ('RIGHTPADDING', (0,0),(-1,-1), 0),
            ]))
            rows.append([sub])
            rows.append([bar])
            rows.append([Spacer(1, 3*mm)])

        t = Table(rows, colWidths=[90*mm])
        t.setStyle(TableStyle([
            ('LEFTPADDING', (0,0),(-1,-1), 6),
            ('RIGHTPADDING', (0,0),(-1,-1), 0),
            ('TOPPADDING', (0,0),(-1,-1), 2),
            ('BOTTOMPADDING', (0,0),(-1,-1), 2),
        ]))
        return t

    # ─── ESRS Completeness ────────────────────────────────────────────────────

    def _make_esrs_completeness(self, by_section: dict) -> list:
        story = []
        story += self._section_header('Complétude par section ESRS', C_NAVY)
        story.append(Paragraph(
            'Taux de remplissage des indicateurs par section ESRS obligatoire.',
            ParagraphStyle('sub', fontSize=9, textColor=C_GRAY, spaceAfter=8)
        ))

        pillar_groups = {
            'environmental': ('Pilier Environnemental (E1–E5)', C_GREEN, C_GREEN2),
            'social':        ('Pilier Social (S1–S4)',          C_BLUE,  C_BLUE2),
            'governance':    ('Pilier Gouvernance (G1)',        C_PURPLE, C_PURPLE2),
        }

        for pillar_key, (pillar_label, p_color, p_bg) in pillar_groups.items():
            sections = [s for s in ESRS_SECTIONS if s['pillar'] == pillar_key]

            # Pillar header
            ph = Table(
                [[Paragraph(f'<font color="white"><b>{pillar_label}</b></font>',
                            ParagraphStyle('ph', fontSize=10))]],
                colWidths=[170*mm],
            )
            ph.setStyle(TableStyle([
                ('BACKGROUND', (0,0),(-1,-1), p_color),
                ('TOPPADDING', (0,0),(-1,-1), 6),
                ('BOTTOMPADDING', (0,0),(-1,-1), 6),
                ('LEFTPADDING', (0,0),(-1,-1), 10),
            ]))
            story.append(ph)

            rows = []
            for sec in sections:
                count = len(by_section.get(sec['code'], []))
                # Heuristic: 5 entries = 100%
                pct = min(round(count / 5 * 100), 100) if count > 0 else 0

                badge = Table(
                    [[Paragraph(
                        f'<font color="white"><b>{sec["code"]}</b></font>',
                        ParagraphStyle('badge', fontSize=8, alignment=1)
                    )]],
                    colWidths=[14*mm], rowHeights=[7*mm],
                )
                badge.setStyle(TableStyle([
                    ('BACKGROUND', (0,0),(-1,-1), p_color),
                    ('ALIGN', (0,0),(-1,-1), 'CENTER'),
                    ('VALIGN', (0,0),(-1,-1), 'MIDDLE'),
                    ('TOPPADDING', (0,0),(-1,-1), 1),
                    ('BOTTOMPADDING', (0,0),(-1,-1), 1),
                    ('ROUNDEDCORNERS', [4]),
                ]))

                bar = self._progress_bar(pct, p_color, width=80*mm, height=5*mm)

                pct_para = Paragraph(
                    f'<font color="{_hex(p_color)}"><b>{pct}%</b></font>',
                    ParagraphStyle('pct', fontSize=9, alignment=2)
                )
                label_para = Paragraph(
                    sec['label'],
                    ParagraphStyle('lbl', fontSize=8, textColor=C_NAVY)
                )
                count_para = Paragraph(
                    f'{count} indicateur(s)',
                    ParagraphStyle('cnt', fontSize=7, textColor=C_GRAY)
                )

                label_col = Table(
                    [[label_para], [count_para]],
                    colWidths=[55*mm],
                )
                label_col.setStyle(TableStyle([
                    ('TOPPADDING', (0,0),(-1,-1), 0),
                    ('BOTTOMPADDING', (0,0),(-1,-1), 1),
                    ('LEFTPADDING', (0,0),(-1,-1), 0),
                    ('RIGHTPADDING', (0,0),(-1,-1), 0),
                ]))

                bar_col = Table(
                    [[bar], [Spacer(1, 1)]],
                    colWidths=[80*mm],
                )
                bar_col.setStyle(TableStyle([
                    ('TOPPADDING', (0,0),(-1,-1), 2),
                    ('BOTTOMPADDING', (0,0),(-1,-1), 0),
                    ('LEFTPADDING', (0,0),(-1,-1), 0),
                    ('RIGHTPADDING', (0,0),(-1,-1), 0),
                ]))

                rows.append([badge, label_col, bar_col, pct_para])

            section_table = Table(
                rows,
                colWidths=[16*mm, 57*mm, 82*mm, 15*mm],
            )
            section_table.setStyle(TableStyle([
                ('ROWBACKGROUNDS', (0,0), (-1,-1), [C_WHITE, p_bg]),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING', (0,0), (-1,-1), 5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                ('LEFTPADDING', (0,0), (-1,-1), 6),
                ('RIGHTPADDING', (0,0), (-1,-1), 4),
                ('GRID', (0,0), (-1,-1), 0.2, C_MGRAY),
            ]))
            story.append(section_table)
            story.append(Spacer(1, 4*mm))

        return story

    # ─── Pillar Data Tables ───────────────────────────────────────────────────

    def _make_pillar_data(self, entries: dict) -> list:
        story = []
        pillars = [
            ('environmental', 'Données Environnementales', C_GREEN),
            ('social',        'Données Sociales',          C_BLUE),
            ('governance',    'Données de Gouvernance',    C_PURPLE),
        ]

        for key, label, color in pillars:
            data = entries.get(key, [])
            if not data:
                continue

            story += self._section_header(label, color)

            rows = [['Métrique', 'Valeur', 'Unité', 'Statut']]
            for entry in data[:40]:
                status = '✓ Vérifié' if entry['verification_status'] == 'verified' else '⏳ Attente'
                rows.append([
                    str(entry['metric_name'] or '')[:50],
                    str(entry['value'] or '-')[:22],
                    str(entry['unit']   or '-')[:12],
                    status,
                ])

            t = Table(rows, colWidths=[80*mm, 35*mm, 25*mm, 30*mm])
            t.setStyle(TableStyle([
                ('BACKGROUND',    (0,0), (-1,0), color),
                ('TEXTCOLOR',     (0,0), (-1,0), C_WHITE),
                ('FONTNAME',      (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE',      (0,0), (-1,0), 9),
                ('FONTSIZE',      (0,1), (-1,-1), 8),
                ('ROWBACKGROUNDS',(0,1), (-1,-1), [C_WHITE, C_LGRAY]),
                ('GRID',          (0,0), (-1,-1), 0.3, C_MGRAY),
                ('ALIGN',         (1,0), (-1,-1), 'CENTER'),
                ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING',    (0,0), (-1,-1), 5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                ('LEFTPADDING',   (0,0), (0,-1), 8),
            ]))
            story.append(t)
            story.append(Paragraph(
                f'Total: {len(data)} indicateur(s)',
                ParagraphStyle('total', fontSize=8, textColor=C_GRAY, spaceAfter=4)
            ))
            story.append(PageBreak())

        return story

    # ─── Compliance Page ──────────────────────────────────────────────────────

    def _make_annexes(self, data: Dict[str, Any]) -> list:
        """Build the annexes section with materiality, risks, suppliers, score history."""
        story = []
        materiality_issues = data.get('materiality_issues', [])
        esg_risks          = data.get('esg_risks', [])
        suppliers          = data.get('suppliers', [])
        score_history      = data.get('score_history', [])

        if not (materiality_issues or esg_risks or suppliers or score_history):
            return story

        # ── Annexe A — Matérialité ──
        if materiality_issues:
            story += self._section_header('Annexe A — Matrice de matérialité (CSRD)', C_PURPLE)
            story.append(Paragraph(
                f'{len(materiality_issues)} enjeux ESG identifiés. La matérialité est évaluée selon une approche '
                'double (impact financier × impact ESG), conformément aux standards CSRD/ESRS. Le graphique '
                'positionne chaque enjeu : les points dans le quadrant supérieur-droit (impact ESG ET financier '
                "élevés) sont jugés matériels au sens CSRD.",
                ParagraphStyle('sub', fontSize=9, textColor=C_GRAY, spaceAfter=10)
            ))
            # Visual quadrant first
            story.append(self._make_materiality_quadrant(materiality_issues))
            story.append(Spacer(1, 6*mm))
            mat_data = [['Enjeu', 'Catégorie', 'Imp. Fin.', 'Imp. ESG', 'Matériel', 'Priorité']]
            for m in materiality_issues[:25]:
                mat_data.append([
                    Paragraph((m.get('name') or '')[:60], ParagraphStyle('mn', fontSize=8)),
                    Paragraph((m.get('category') or '—')[:25], ParagraphStyle('mc', fontSize=8)),
                    f"{m.get('financial_impact', 0):.0f}",
                    f"{m.get('esg_impact', 0):.0f}",
                    '✓' if m.get('is_material') else '—',
                    (m.get('priority') or '—')[:12],
                ])
            mt = Table(mat_data, colWidths=[60*mm, 30*mm, 18*mm, 18*mm, 16*mm, 22*mm])
            mt.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), C_PURPLE),
                ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
                ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',   (0, 0), (-1, 0), 9),
                ('ALIGN',      (2, 0), (-1, -1), 'CENTER'),
                ('VALIGN',     (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID',       (0, 0), (-1, -1), 0.3, colors.HexColor('#E5E7EB')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
                ('FONTSIZE',   (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(mt)
            if len(materiality_issues) > 25:
                story.append(Paragraph(
                    f'<i>… et {len(materiality_issues) - 25} autres enjeux. Tableau complet dans l\'export Excel.</i>',
                    ParagraphStyle('more', fontSize=8, textColor=C_GRAY, spaceBefore=6),
                ))
            story.append(PageBreak())

        # ── Annexe B — Risques ESG ──
        if esg_risks:
            story += self._section_header('Annexe B — Registre des risques ESG', C_RED)
            story.append(Paragraph(
                f'{len(esg_risks)} risques ESG identifiés et évalués sur une échelle 1-5 (probabilité × impact). '
                'Score ≥ 20 = critique, 12-19 = élevé, 6-11 = modéré, < 6 = faible.',
                ParagraphStyle('sub', fontSize=9, textColor=C_GRAY, spaceAfter=10)
            ))
            # Visual heatmap first
            story.append(self._make_risk_heatmap(esg_risks))
            story.append(Spacer(1, 6*mm))
            risk_data = [['Risque', 'Catégorie', 'Prob.', 'Imp.', 'Score', 'Mitigation']]
            for r in esg_risks[:25]:
                score = r.get('risk_score') or (int(r.get('probability', 0)) * int(r.get('impact', 0)))
                risk_data.append([
                    Paragraph((r.get('title') or '')[:60], ParagraphStyle('rn', fontSize=8)),
                    Paragraph((r.get('category') or '—')[:25], ParagraphStyle('rc', fontSize=8)),
                    str(r.get('probability', 0)),
                    str(r.get('impact', 0)),
                    str(score),
                    Paragraph((r.get('mitigation_status') or '—')[:30], ParagraphStyle('rm', fontSize=8)),
                ])
            rt = Table(risk_data, colWidths=[60*mm, 28*mm, 14*mm, 14*mm, 14*mm, 34*mm])
            rt.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), C_RED),
                ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
                ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',   (0, 0), (-1, 0), 9),
                ('ALIGN',      (2, 0), (4, -1), 'CENTER'),
                ('VALIGN',     (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID',       (0, 0), (-1, -1), 0.3, colors.HexColor('#E5E7EB')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
                ('FONTSIZE',   (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(rt)
            if len(esg_risks) > 25:
                story.append(Paragraph(
                    f'<i>… et {len(esg_risks) - 25} autres risques. Tableau complet dans l\'export Excel.</i>',
                    ParagraphStyle('more', fontSize=8, textColor=C_GRAY, spaceBefore=6),
                ))
            story.append(PageBreak())

        # ── Annexe C — Fournisseurs ──
        if suppliers:
            story += self._section_header('Annexe C — Chaîne d\'approvisionnement (top fournisseurs)', C_BLUE)
            story.append(Paragraph(
                f'Évaluation ESG de {min(len(suppliers), 20)} fournisseurs sur un total de {len(suppliers)}. '
                'Tri décroissant par score global ESG.',
                ParagraphStyle('sub', fontSize=9, textColor=C_GRAY, spaceAfter=10)
            ))
            # Visual ranking bar chart first
            story.append(self._make_supplier_top_chart(suppliers))
            story.append(Spacer(1, 6*mm))
            sup_data = [['Fournisseur', 'Pays', 'Risque', 'Score Global', 'E', 'S', 'G']]
            for s in suppliers[:20]:
                sup_data.append([
                    Paragraph((s.get('name') or '')[:50], ParagraphStyle('sn', fontSize=8)),
                    (s.get('country') or '—')[:15],
                    s.get('risk_level', '—'),
                    f"{s.get('global_score'):.0f}" if s.get('global_score') is not None else '—',
                    f"{s.get('env_score'):.0f}"   if s.get('env_score')   is not None else '—',
                    f"{s.get('social_score'):.0f}" if s.get('social_score') is not None else '—',
                    f"{s.get('gov_score'):.0f}"   if s.get('gov_score')   is not None else '—',
                ])
            st = Table(sup_data, colWidths=[55*mm, 22*mm, 18*mm, 22*mm, 14*mm, 14*mm, 14*mm])
            st.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), C_BLUE),
                ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
                ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',   (0, 0), (-1, 0), 9),
                ('ALIGN',      (1, 0), (-1, -1), 'CENTER'),
                ('VALIGN',     (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID',       (0, 0), (-1, -1), 0.3, colors.HexColor('#E5E7EB')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
                ('FONTSIZE',   (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(st)
            story.append(PageBreak())

        # ── Annexe D — Historique des scores ──
        if score_history:
            story += self._section_header('Annexe D — Historique du score ESG', C_GREEN)
            story.append(Paragraph(
                f'Évolution du score sur les {len(score_history)} dernières mesures. La courbe ci-dessous '
                'présente la trajectoire du score global et par pilier (E/S/G).',
                ParagraphStyle('sub', fontSize=9, textColor=C_GRAY, spaceAfter=10)
            ))
            # Visual chart first
            story.append(self._make_score_line_chart(score_history))
            story.append(Spacer(1, 6*mm))
            hist_data = [['Date', 'Global', 'Note', 'Env.', 'Soc.', 'Gouv.', 'Complétude']]
            for s in score_history:
                hist_data.append([
                    s.get('date') or '—',
                    f"{s.get('overall', 0):.1f}",
                    s.get('rating') or '—',
                    f"{s.get('environmental', 0):.1f}",
                    f"{s.get('social', 0):.1f}",
                    f"{s.get('governance', 0):.1f}",
                    f"{s.get('data_completeness', 0):.0f}%",
                ])
            ht = Table(hist_data, colWidths=[28*mm, 22*mm, 18*mm, 22*mm, 22*mm, 22*mm, 24*mm])
            ht.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), C_GREEN),
                ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
                ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',   (0, 0), (-1, 0), 9),
                ('ALIGN',      (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN',     (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID',       (0, 0), (-1, -1), 0.3, colors.HexColor('#E5E7EB')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
                ('FONTSIZE',   (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(ht)
            story.append(PageBreak())

        return story

    def _make_compliance_page(self, year: int) -> list:
        story = []
        story += self._section_header('Conformité & Cadres de référence', C_NAVY)
        story.append(Paragraph(
            'Ce rapport est conforme aux standards internationaux de reporting ESG en vigueur.',
            ParagraphStyle('sub', fontSize=9, textColor=C_GRAY, spaceAfter=10)
        ))

        frameworks = [
            ('CSRD', 'Corporate Sustainability Reporting Directive', C_GREEN,
             'Directive européenne 2022/2464 — Obligatoire à partir de 2024'),
            ('ESRS', 'European Sustainability Reporting Standards', C_GREEN,
             'Standards EFRAG — 12 normes thématiques couvrant E, S et G'),
            ('GRI', 'Global Reporting Initiative', C_BLUE,
             'GRI Standards 2021 — Référentiel mondial de durabilité'),
            ('TCFD', 'Task Force on Climate-related Financial Disclosures', C_BLUE,
             'Recommandations sur les risques et opportunités climatiques'),
            ('GHG Protocol', 'Greenhouse Gas Protocol', C_PURPLE,
             'Comptabilisation et reporting des émissions de GES (scopes 1, 2, 3)'),
            ('ISO 26000', 'Norme ISO de responsabilité sociétale', C_PURPLE,
             'Lignes directrices pour la responsabilité des organisations'),
        ]

        rows = []
        for acronym, name, color, desc in frameworks:
            badge = Table(
                [[Paragraph(
                    f'<font color="white"><b>{acronym}</b></font>',
                    ParagraphStyle('fb', fontSize=9, alignment=1)
                )]],
                colWidths=[22*mm], rowHeights=[9*mm],
            )
            badge.setStyle(TableStyle([
                ('BACKGROUND', (0,0),(-1,-1), color),
                ('ALIGN', (0,0),(-1,-1), 'CENTER'),
                ('VALIGN', (0,0),(-1,-1), 'MIDDLE'),
                ('ROUNDEDCORNERS', [4]),
            ]))
            rows.append([
                badge,
                Table(
                    [[Paragraph(f'<b>{name}</b>',
                                ParagraphStyle('fn', fontSize=9, textColor=C_NAVY))],
                     [Paragraph(desc,
                                ParagraphStyle('fd', fontSize=8, textColor=C_GRAY))]],
                    colWidths=[145*mm],
                )
            ])

        t = Table(rows, colWidths=[24*mm, 146*mm])
        t.setStyle(TableStyle([
            ('ROWBACKGROUNDS', (0,0), (-1,-1), [C_WHITE, C_LGRAY]),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('GRID', (0,0), (-1,-1), 0.3, C_MGRAY),
        ]))
        story.append(t)
        story.append(Spacer(1, 10*mm))

        story.append(HRFlowable(width='100%', thickness=1, color=C_MGRAY))
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph(
            f'© {year} ESG Flow · Tous droits réservés · Rapport généré automatiquement par la plateforme ESG Flow v2.0.0',
            ParagraphStyle('footer', fontSize=8, textColor=C_GRAY, alignment=1)
        ))
        return story

    # ─── Helpers ──────────────────────────────────────────────────────────────

    def _section_header(self, title: str, color, size: int = 15) -> list:
        return [
            Paragraph(
                title,
                ParagraphStyle('sh', fontSize=size, fontName='Helvetica-Bold',
                               textColor=color, spaceBefore=6, spaceAfter=4)
            ),
            HRFlowable(width='100%', thickness=1.5, color=color, spaceAfter=8),
        ]

    def _make_calculated_kpis_section(self, calculated_kpis: dict) -> list:
        """KPIs Automatiques ESRS — section ReportLab avec les formules calculées."""
        story = self._section_header('KPIs Calculés Automatiquement', C_GREEN)
        story.append(Paragraph(
            "Métriques dérivées automatiquement des données saisies — "
            "conformes aux exigences ESRS E1-6, E1-4, S1-8, S1-13 et S1-16.",
            ParagraphStyle('kpi_sub', fontSize=8, textColor=C_GRAY, spaceAfter=10, leading=12)
        ))

        if not calculated_kpis:
            story.append(Paragraph(
                "⚠  Aucune métrique calculée disponible. Saisissez vos données dans "
                "Calculs Automatiques (Données › Calcul auto) pour activer les formules ESRS.",
                ParagraphStyle('kpi_empty', fontSize=9, textColor=C_AMBER,
                               spaceAfter=12, leading=13, borderPad=8,
                               backColor=C_AMBER2, borderWidth=1, borderColor=C_AMBER,
                               borderRadius=4)
            ))
            return story

        ESRS_REF = {
            'scope_3_total':        'ESRS E1-6',
            'carbon_intensity':     'ESRS E1-6',
            'renewable_percentage': 'ESRS E1-4',
            'turnover_rate':        'ESRS S1-8',
            'training_rate':        'ESRS S1-13',
            'women_percentage':     'ESRS S1-16',
        }

        PILLAR_HEX = {
            'environmental': '#059669',
            'social':        '#2563eb',
            'governance':    '#7c3aed',
        }
        PILLAR_COLOR = {
            'environmental': C_GREEN,
            'social':        C_BLUE,
            'governance':    C_PURPLE,
        }

        # Table header
        rows = [[
            Paragraph('<b>KPI ESRS</b>',   ParagraphStyle('kh', fontSize=9, fontName='Helvetica-Bold', textColor=C_WHITE)),
            Paragraph('<b>Résultat</b>',   ParagraphStyle('kh', fontSize=9, fontName='Helvetica-Bold', textColor=C_WHITE)),
            Paragraph('<b>Inputs utilisés</b>', ParagraphStyle('kh', fontSize=9, fontName='Helvetica-Bold', textColor=C_WHITE)),
            Paragraph('<b>Norme</b>',       ParagraphStyle('kh', fontSize=9, fontName='Helvetica-Bold', textColor=C_WHITE)),
        ]]

        for key, kpi in calculated_kpis.items():
            name    = kpi.get('name') or key.replace('_', ' ').title()
            value   = kpi.get('value') or 0
            unit    = kpi.get('unit', '')
            inputs  = kpi.get('inputs_used', {})
            pillar  = kpi.get('pillar', 'environmental')
            col_hex = PILLAR_HEX.get(pillar, '#059669')
            esrs    = ESRS_REF.get(key, 'ESRS')

            inputs_lines = '<br/>'.join(
                f"{k.replace('_', ' ').capitalize()}: <b>{v:,.2f}</b>"
                for k, v in inputs.items()
            )

            rows.append([
                Paragraph(f'<b>{name}</b>', ParagraphStyle('kn', fontSize=9, textColor=C_NAVY, leading=12)),
                Paragraph(
                    f'<font color="{col_hex}"><b>{value:,.2f}</b></font><br/>'
                    f'<font color="#6b7280" size="8">{unit}</font>',
                    ParagraphStyle('kv', fontSize=10, leading=13)
                ),
                Paragraph(inputs_lines or '—', ParagraphStyle('ki', fontSize=8, textColor=C_GRAY, leading=11)),
                Paragraph(
                    f'<font color="#2563eb"><b>{esrs}</b></font>',
                    ParagraphStyle('kr', fontSize=8, leading=11)
                ),
            ])

        t = Table(rows, colWidths=[60*mm, 32*mm, 62*mm, 21*mm])
        t.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, 0),  C_GREEN),
            ('ROWBACKGROUNDS',(0, 1), (-1, -1), [C_WHITE, C_LGRAY]),
            ('TOPPADDING',    (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 9),
            ('LEFTPADDING',   (0, 0), (-1, -1), 8),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 8),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID',          (0, 0), (-1, -1), 0.4, colors.HexColor('#e5e7eb')),
            ('LINEBELOW',     (0, -1), (-1, -1), 2, C_GREEN),
        ]))
        story.append(t)
        story.append(Spacer(1, 5*mm))
        story.append(Paragraph(
            "<i>Ces KPIs sont calculés automatiquement depuis les données saisies. "
            "L'exactitude dépend de la complétude des entrées ESG.</i>",
            ParagraphStyle('kpi_note', fontSize=7, textColor=C_GRAY, spaceAfter=6, leading=10)
        ))
        return story

    def _progress_bar(self, pct: int, color, width=90*mm, height=5*mm) -> Drawing:
        w = float(width)
        h = float(height)
        d = Drawing(w, h)
        # Background track
        d.add(Rect(0, 0, w, h, rx=2, ry=2,
                   fillColor=C_MGRAY, strokeColor=None))
        # Fill
        fill_w = max(w * pct / 100, 0.1)
        d.add(Rect(0, 0, fill_w, h, rx=2, ry=2,
                   fillColor=color, strokeColor=None))
        return d

    # ─── Modern charts (native ReportLab — no extra deps) ─────────────────────

    def _make_score_line_chart(
        self, history: list, width=170*mm, height=55*mm
    ) -> Drawing:
        """Line chart of ESG score evolution.
        history: list of dicts with 'date', 'overall', 'environmental', 'social', 'governance'.
        Renders multi-series line plot with legend.
        """
        w, h = float(width), float(height)
        d = Drawing(w, h)
        if not history or len(history) < 2:
            d.add(String(w/2, h/2, 'Pas assez de données historiques pour tracer la courbe.',
                         fontSize=8, fillColor=C_GRAY, textAnchor='middle'))
            return d

        # Most recent points come last
        ordered = list(reversed(history))[-12:]  # last 12 points
        labels = [(p.get('date') or '')[2:7].replace('-', '/') for p in ordered]  # YY/MM
        def _series(key):
            return [(i, float(p.get(key) or 0)) for i, p in enumerate(ordered)]

        lp = LinePlot()
        lp.x = 36
        lp.y = 28
        lp.width  = w - 50
        lp.height = h - 50
        lp.data = [_series('overall'), _series('environmental'),
                   _series('social'),  _series('governance')]
        # Style each line
        series_styles = [
            (C_NAVY,   2.4),
            (C_GREEN,  1.4),
            (C_BLUE,   1.4),
            (C_PURPLE, 1.4),
        ]
        for i, (color, w_) in enumerate(series_styles):
            lp.lines[i].strokeColor = color
            lp.lines[i].strokeWidth = w_
            lp.lines[i].symbol = makeMarker('Circle')
            lp.lines[i].symbol.fillColor = color
            lp.lines[i].symbol.strokeColor = colors.white
            lp.lines[i].symbol.strokeWidth = 1
            lp.lines[i].symbol.size = 4

        # Axes
        lp.xValueAxis.valueMin = 0
        lp.xValueAxis.valueMax = len(ordered) - 1
        lp.xValueAxis.valueSteps = list(range(len(ordered)))
        lp.xValueAxis.labelTextFormat = lambda v: labels[int(v)] if 0 <= int(v) < len(labels) else ''
        lp.xValueAxis.labels.fontSize = 7
        lp.xValueAxis.labels.fillColor = C_GRAY
        lp.xValueAxis.strokeColor = C_MGRAY
        lp.xValueAxis.visibleTicks = False
        lp.xValueAxis.labels.angle = 0

        lp.yValueAxis.valueMin = 0
        lp.yValueAxis.valueMax = 100
        lp.yValueAxis.valueStep = 20
        lp.yValueAxis.labels.fontSize = 7
        lp.yValueAxis.labels.fillColor = C_GRAY
        lp.yValueAxis.strokeColor = C_MGRAY
        lp.yValueAxis.visibleTicks = False
        lp.yValueAxis.visibleGrid = True
        lp.yValueAxis.gridStrokeColor = colors.HexColor('#f1f5f9')
        lp.yValueAxis.gridStrokeWidth = 0.5

        d.add(lp)

        # Legend
        legend_x = 36
        legend_y = h - 8
        items = [
            ('Global',         C_NAVY),
            ('Environnement',  C_GREEN),
            ('Social',         C_BLUE),
            ('Gouvernance',    C_PURPLE),
        ]
        for label, color in items:
            d.add(Rect(legend_x, legend_y, 8, 3, fillColor=color, strokeColor=None))
            d.add(String(legend_x + 12, legend_y, label,
                         fontSize=7, fillColor=C_NAVY))
            legend_x += 42
        return d

    def _make_scopes_donut(
        self, scope1: float, scope2: float, scope3: float,
        width=80*mm, height=70*mm
    ) -> Drawing:
        """Donut chart: emissions split by GHG scope (1/2/3 in tCO2e)."""
        w, h = float(width), float(height)
        d = Drawing(w, h)
        total = scope1 + scope2 + scope3
        if total <= 0:
            d.add(String(w/2, h/2, 'Aucune émission calculée.',
                         fontSize=8, fillColor=C_GRAY, textAnchor='middle'))
            return d
        cx, cy, r = w/2, h/2 + 6, 26
        pie = Pie()
        pie.x = cx - r
        pie.y = cy - r
        pie.width = pie.height = 2 * r
        pie.data = [max(scope1, 0.01), max(scope2, 0.01), max(scope3, 0.01)]
        pie.innerRadiusFraction = 0.55
        pie.simpleLabels = 1
        pie.labels = ['', '', '']
        cols = [C_GREEN, C_BLUE, C_PURPLE]
        for i, c in enumerate(cols):
            pie.slices[i].fillColor = c
            pie.slices[i].strokeColor = colors.white
            pie.slices[i].strokeWidth = 1.5
        d.add(pie)
        # Center total
        d.add(String(cx, cy + 3, f'{total:,.1f}',
                     fontSize=12, fontName='Helvetica-Bold',
                     fillColor=C_NAVY, textAnchor='middle'))
        d.add(String(cx, cy - 8, 'tCO₂e',
                     fontSize=7, fillColor=C_GRAY, textAnchor='middle'))
        # Legend rows under donut
        y0 = 10
        for i, (label, val, color) in enumerate([
            (f'Scope 1', scope1, C_GREEN),
            (f'Scope 2', scope2, C_BLUE),
            (f'Scope 3', scope3, C_PURPLE),
        ]):
            x = 6 + i * (w - 12) / 3
            d.add(Rect(x, y0, 6, 6, fillColor=color, strokeColor=None))
            pct = (val / total * 100) if total else 0
            d.add(String(x + 9, y0 + 1, f'{label} · {pct:.0f}%',
                         fontSize=7, fillColor=C_NAVY))
        return d

    def _make_pillar_bar_chart(
        self, env: int, soc: int, gov: int,
        width=80*mm, height=50*mm
    ) -> Drawing:
        """Vertical bar chart for indicator counts per pillar (cover/exec summary)."""
        w, h = float(width), float(height)
        d = Drawing(w, h)
        bc = VerticalBarChart()
        bc.x = 32
        bc.y = 18
        bc.width  = w - 40
        bc.height = h - 30
        bc.data = [[env, soc, gov]]
        bc.barWidth = 12
        bc.groupSpacing = 14
        bc.bars[0].fillColor = C_GREEN
        bc.bars.strokeColor = None
        bc.valueAxis.valueMin = 0
        max_v = max(env, soc, gov, 1)
        bc.valueAxis.valueMax = max_v * 1.2
        bc.valueAxis.valueStep = max(1, int(max_v / 4))
        bc.valueAxis.labels.fontSize = 7
        bc.valueAxis.labels.fillColor = C_GRAY
        bc.valueAxis.strokeColor = C_MGRAY
        bc.valueAxis.visibleGrid = True
        bc.valueAxis.gridStrokeColor = colors.HexColor('#f1f5f9')
        bc.valueAxis.gridStrokeWidth = 0.5
        bc.categoryAxis.categoryNames = ['Env.', 'Social', 'Gouv.']
        bc.categoryAxis.labels.fontSize = 8
        bc.categoryAxis.labels.fillColor = C_NAVY
        bc.categoryAxis.strokeColor = C_MGRAY
        # Custom per-bar coloring via patches (one series, 3 bars)
        for i, color in enumerate([C_GREEN, C_BLUE, C_PURPLE]):
            try:
                bc.bars[(0, i)].fillColor = color
            except Exception:
                pass
        d.add(bc)
        d.add(String(w/2, h - 6, 'Indicateurs par pilier',
                     fontSize=8, fontName='Helvetica-Bold',
                     fillColor=C_NAVY, textAnchor='middle'))
        return d

    def _make_horizontal_bar_chart(
        self, items: list, color=None, max_val: float = None,
        width=170*mm, height=None
    ) -> Drawing:
        """Horizontal bar chart for ranked comparisons (top sectors, top suppliers).
        items: list of (label, value) tuples, sorted desc.
        """
        if not items:
            d = Drawing(float(width), 10*mm)
            return d
        items = items[:8]  # cap
        bar_h = 6.5*mm
        gap = 2*mm
        h = float(height) if height else (len(items) * (bar_h + gap) + 14)
        w = float(width)
        d = Drawing(w, h)
        max_v = max_val if max_val is not None else max((v for _, v in items if v is not None), default=1)
        if max_v <= 0:
            max_v = 1
        c = color or C_GREEN
        chart_x = 70*mm     # leave room for labels
        chart_w = w - chart_x - 24*mm
        for i, (label, val) in enumerate(items):
            y = h - (i + 1) * (bar_h + gap)
            val_f = float(val or 0)
            bar_w = max(chart_w * val_f / max_v, 0.5)
            # Label (left)
            d.add(String(chart_x - 4, y + 2, str(label)[:35],
                         fontSize=8, fillColor=C_NAVY, textAnchor='end'))
            # Track
            d.add(Rect(chart_x, y, chart_w, bar_h, rx=1.5, ry=1.5,
                       fillColor=colors.HexColor('#f1f5f9'), strokeColor=None))
            # Bar
            d.add(Rect(chart_x, y, bar_w, bar_h, rx=1.5, ry=1.5,
                       fillColor=c, strokeColor=None))
            # Value
            d.add(String(chart_x + chart_w + 4, y + 2,
                         f'{val_f:,.1f}' if isinstance(val, float) else str(val),
                         fontSize=8, fillColor=C_NAVY))
        return d

    def _make_materiality_quadrant(
        self, issues: list, width=130*mm, height=100*mm
    ) -> Drawing:
        """Double-materiality matrix (financial impact vs ESG impact).
        Each issue is a dot; quadrant labels indicate materiality status.
        """
        w, h = float(width), float(height)
        d = Drawing(w, h)
        ml, mb = 24, 22
        mr, mt = 8, 14
        chart_w = w - ml - mr
        chart_h = h - mb - mt
        # Background quadrants
        d.add(Rect(ml, mb, chart_w/2, chart_h/2,
                   fillColor=colors.HexColor('#f8fafc'), strokeColor=None))  # low/low
        d.add(Rect(ml + chart_w/2, mb, chart_w/2, chart_h/2,
                   fillColor=colors.HexColor('#fef3c7'), strokeColor=None))  # fin only
        d.add(Rect(ml, mb + chart_h/2, chart_w/2, chart_h/2,
                   fillColor=colors.HexColor('#d1fae5'), strokeColor=None))  # esg only
        d.add(Rect(ml + chart_w/2, mb + chart_h/2, chart_w/2, chart_h/2,
                   fillColor=colors.HexColor('#bfdbfe'), strokeColor=None))  # both (material)
        # Grid borders
        d.add(Line(ml, mb, ml + chart_w, mb, strokeColor=C_MGRAY, strokeWidth=0.6))
        d.add(Line(ml, mb, ml, mb + chart_h, strokeColor=C_MGRAY, strokeWidth=0.6))
        d.add(Line(ml, mb + chart_h/2, ml + chart_w, mb + chart_h/2,
                   strokeColor=C_MGRAY, strokeWidth=0.4, strokeDashArray=[2, 2]))
        d.add(Line(ml + chart_w/2, mb, ml + chart_w/2, mb + chart_h,
                   strokeColor=C_MGRAY, strokeWidth=0.4, strokeDashArray=[2, 2]))
        # Axis labels
        d.add(String(ml + chart_w/2, 6, 'Impact financier (0-100) →',
                     fontSize=7, fillColor=C_GRAY, textAnchor='middle'))
        d.add(String(8, mb + chart_h/2, 'Impact ESG ↑',
                     fontSize=7, fillColor=C_GRAY, textAnchor='middle'))
        # Quadrant labels
        d.add(String(ml + chart_w*0.75, mb + chart_h - 6, 'MATÉRIEL',
                     fontSize=7, fontName='Helvetica-Bold',
                     fillColor=C_BLUE, textAnchor='middle'))
        # Plot dots
        for it in issues[:50]:
            fx = float(it.get('financial_impact') or 0)
            ex = float(it.get('esg_impact') or 0)
            x = ml + (fx / 100) * chart_w
            y = mb + (ex / 100) * chart_h
            is_mat = bool(it.get('is_material'))
            color = C_NAVY if is_mat else C_GRAY
            d.add(Circle(x, y, 3.5 if is_mat else 2.5,
                         fillColor=color, strokeColor=colors.white, strokeWidth=1))
        # Title
        d.add(String(w/2, h - 6, 'Matrice de double matérialité',
                     fontSize=9, fontName='Helvetica-Bold',
                     fillColor=C_NAVY, textAnchor='middle'))
        return d

    def _make_insight_card(self, title: str, message: str, color, bg_color, icon: str = '◆') -> Table:
        """Card with a colored stripe, title, and short narrative — for executive insights."""
        body = Table(
            [[Paragraph(f'<font color="{_hex(color)}"><b>{icon}  {title}</b></font>',
                        ParagraphStyle('ict', fontSize=9, leading=12, spaceAfter=2)),
              ],
             [Paragraph(message,
                        ParagraphStyle('icm', fontSize=8.5, textColor=C_NAVY,
                                       leading=12))]],
            colWidths=[160*mm],
        )
        body.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), bg_color),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (0, 0), 7),
            ('BOTTOMPADDING', (0, -1), (-1, -1), 7),
            ('TOPPADDING', (0, 1), (0, 1), 1),
            ('BOTTOMPADDING', (0, 0), (0, 0), 2),
            ('LINEBEFORE', (0, 0), (0, -1), 3, color),
            ('ROUNDEDCORNERS', [4]),
        ]))
        return body

    # GHG Protocol Scope 3 — 15 upstream/downstream categories
    SCOPE3_CATEGORIES = [
        (1,  'Achats de biens et services'),
        (2,  'Biens d\'équipement'),
        (3,  'Activités liées aux combustibles et à l\'énergie'),
        (4,  'Transport et distribution amont'),
        (5,  'Déchets générés'),
        (6,  'Déplacements professionnels'),
        (7,  'Trajets domicile-travail'),
        (8,  'Actifs loués en amont'),
        (9,  'Transport et distribution aval'),
        (10, 'Transformation des produits vendus'),
        (11, 'Utilisation des produits vendus'),
        (12, 'Fin de vie des produits vendus'),
        (13, 'Actifs loués en aval'),
        (14, 'Franchises'),
        (15, 'Investissements'),
    ]

    def _aggregate_scopes(self, entries: dict) -> Dict[str, Any]:
        """Aggregate emissions by GHG scope (1/2/3) from environmental entries.
        Returns dict with scope1/scope2/scope3 in tCO2e + scope3_categories detail.
        Heuristic: looks for 'scope' keywords in metric name/category and
        normalises numeric values to tonnes (kg → t).
        """
        env_entries = entries.get('environmental', []) if isinstance(entries, dict) else []
        out: Dict[str, Any] = {'scope1': 0.0, 'scope2': 0.0, 'scope3': 0.0, 'scope3_categories': {}}
        for e in env_entries:
            text = (str(e.get('metric_name') or '') + ' ' + str(e.get('category') or '')).lower()
            unit = str(e.get('unit') or '').lower().strip()
            try:
                val = float(e.get('value') or 0)
            except (TypeError, ValueError):
                continue
            if val <= 0:
                continue
            # Normalise to tonnes
            if 'kg' in unit and 't' not in unit.replace('kg', ''):
                val = val / 1000.0
            # Only count CO2e-like values
            if not any(x in (text + ' ' + unit) for x in ('co2', 'ges', 'ghg', 'tco2', 'kgco2', 'carbone', 'émission', 'emission')):
                continue
            if 'scope 1' in text or 'scope1' in text:
                out['scope1'] += val
            elif 'scope 2' in text or 'scope2' in text:
                out['scope2'] += val
            elif 'scope 3' in text or 'scope3' in text:
                out['scope3'] += val
                # Try to identify Scope 3 sub-category (cat 1-15)
                for cat_num, cat_label in self.SCOPE3_CATEGORIES:
                    aliases = [f'cat {cat_num}', f'cat{cat_num}', f'catégorie {cat_num}',
                               f'categorie {cat_num}', f'category {cat_num}',
                               cat_label.lower()[:20]]
                    if any(a in text for a in aliases):
                        out['scope3_categories'][cat_num] = out['scope3_categories'].get(cat_num, 0.0) + val
                        break
        return out

    # ─── New visual helpers — hero score, sparkline, ESG donuts ──────────────

    def _make_score_gauge(
        self, score: float, rating: str = '', width=70*mm, height=70*mm,
    ) -> Drawing:
        """Circular gauge showing the overall ESG score with a colored arc.

        Designed for the cover page — a single, impactful "hero" number.
        The arc length is proportional to the score (0-100); the colour
        ramps from red → amber → blue → green as the score improves.
        """
        from math import cos, sin, radians
        w, h = float(width), float(height)
        d = Drawing(w, h)
        cx, cy = w / 2, h / 2 - 2
        r_outer, r_inner = min(w, h) / 2 - 6, min(w, h) / 2 - 14

        score = max(0.0, min(100.0, float(score or 0)))
        color = (
            C_GREEN if score >= 75 else
            C_BLUE  if score >= 55 else
            C_AMBER if score >= 35 else
            C_RED
        )

        # Background ring (full circle, light gray)
        steps = 60
        bg_pts = []
        for i in range(steps + 1):
            ang = radians(90 - (i / steps) * 360)
            bg_pts += [cx + r_outer * cos(ang), cy + r_outer * sin(ang)]
        d.add(PolyLine(bg_pts, strokeColor=C_LGRAY, strokeWidth=8))

        # Foreground arc (proportional to score)
        fg_pts = []
        arc_steps = max(1, int(steps * score / 100))
        for i in range(arc_steps + 1):
            ang = radians(90 - (i / steps) * 360)
            fg_pts += [cx + r_outer * cos(ang), cy + r_outer * sin(ang)]
        if len(fg_pts) >= 4:
            d.add(PolyLine(fg_pts, strokeColor=color, strokeWidth=8))

        # Centre: big number + "/100" + rating
        d.add(String(cx, cy + 2, f'{score:.0f}',
                     fontSize=30, fontName='Helvetica-Bold',
                     fillColor=color, textAnchor='middle'))
        d.add(String(cx, cy - 10, '/ 100',
                     fontSize=8, fillColor=C_GRAY, textAnchor='middle'))
        if rating and rating != 'N/A':
            d.add(String(cx, cy - 22, f'Note  {rating}',
                         fontSize=8, fontName='Helvetica-Bold',
                         fillColor=C_NAVY, textAnchor='middle'))
        d.add(String(cx, 6, 'SCORE ESG GLOBAL',
                     fontSize=6.5, fontName='Helvetica-Bold',
                     fillColor=C_GRAY, textAnchor='middle'))
        return d

    def _make_score_sparkline(
        self, history: list, width=80*mm, height=22*mm,
    ) -> Drawing:
        """Tiny inline trend chart of the overall score over time.

        Renders a smooth line with a soft gradient fill underneath, plus
        the first/last value annotations. Designed to sit beside or below
        the hero gauge on the cover.
        """
        w, h = float(width), float(height)
        d = Drawing(w, h)
        if not history:
            return d
        # Newest last
        scores = []
        for s in history:
            try:
                scores.append(float(s.get('overall') or 0))
            except (TypeError, ValueError):
                continue
        if not scores:
            return d

        ml, mr, mb, mt = 18, 22, 6, 12
        plot_w = w - ml - mr
        plot_h = h - mb - mt
        d.add(String(w/2, h - 6, 'Trajectoire du score (12 mois)',
                     fontSize=7, fontName='Helvetica-Bold',
                     fillColor=C_GRAY, textAnchor='middle'))

        s_min = min(scores)
        s_max = max(scores)
        if s_max - s_min < 5:  # avoid a flat line
            s_min, s_max = max(0, s_min - 5), min(100, s_max + 5)
        span = max(1.0, s_max - s_min)

        n = len(scores)
        pts = []
        for i, sc in enumerate(scores):
            x = ml + (i / max(1, n - 1)) * plot_w
            y = mb + ((sc - s_min) / span) * plot_h
            pts += [x, y]

        # Filled polygon (sparkline area)
        if len(pts) >= 4:
            area = [pts[0], mb] + pts + [pts[-2], mb]
            from reportlab.graphics.shapes import Polygon
            d.add(Polygon(area, fillColor=colors.HexColor('#d1fae5'),
                          strokeColor=None))
            d.add(PolyLine(pts, strokeColor=C_GREEN, strokeWidth=1.5))

        # First/last value labels
        d.add(String(ml - 2, pts[1] - 3, f'{scores[0]:.0f}',
                     fontSize=6.5, fillColor=C_GRAY, textAnchor='end'))
        d.add(String(ml + plot_w + 2, pts[-1] - 3, f'{scores[-1]:.0f}',
                     fontSize=7, fontName='Helvetica-Bold',
                     fillColor=C_GREEN, textAnchor='start'))
        return d

    def _make_pillar_donuts(
        self, env_score: float, soc_score: float, gov_score: float,
        width=170*mm, height=55*mm,
    ) -> Drawing:
        """Three side-by-side donut charts — one per pillar (E/S/G).

        Each donut shows the pillar's score 0-100; the unfilled portion is
        a light grey ring so the visual mass communicates "completeness".
        """
        from math import cos, sin, radians
        w, h = float(width), float(height)
        d = Drawing(w, h)
        third = w / 3

        pillars = [
            ('Environnement', env_score, C_GREEN,  C_GREEN2),
            ('Social',        soc_score, C_BLUE,   C_BLUE2),
            ('Gouvernance',   gov_score, C_PURPLE, C_PURPLE2),
        ]
        for idx, (label, value, color, bg) in enumerate(pillars):
            cx = third * idx + third / 2
            cy = h / 2 - 2
            r  = min(third * 0.36, h * 0.35)
            value = max(0.0, min(100.0, float(value or 0)))

            # Background full ring (lighter shade of pillar color)
            bg_pts = []
            steps = 48
            for i in range(steps + 1):
                ang = radians(90 - (i / steps) * 360)
                bg_pts += [cx + r * cos(ang), cy + r * sin(ang)]
            d.add(PolyLine(bg_pts, strokeColor=bg, strokeWidth=7))

            # Foreground arc proportional to value
            fg_pts = []
            arc_steps = max(1, int(steps * value / 100))
            for i in range(arc_steps + 1):
                ang = radians(90 - (i / steps) * 360)
                fg_pts += [cx + r * cos(ang), cy + r * sin(ang)]
            if len(fg_pts) >= 4:
                d.add(PolyLine(fg_pts, strokeColor=color, strokeWidth=7))

            d.add(String(cx, cy + 2, f'{value:.0f}',
                         fontSize=20, fontName='Helvetica-Bold',
                         fillColor=color, textAnchor='middle'))
            d.add(String(cx, cy - 10, '/ 100',
                         fontSize=7, fillColor=C_GRAY, textAnchor='middle'))
            d.add(String(cx, 6, label.upper(),
                         fontSize=8, fontName='Helvetica-Bold',
                         fillColor=C_NAVY, textAnchor='middle'))
        return d

    def _make_completeness_bars(
        self, by_section: dict, width=170*mm, height=58*mm,
    ) -> Drawing:
        """ESRS completeness bars per standard (E1..G1).

        Each row is a horizontal progress bar of the % of mandatory data
        points covered for that ESRS section, colored by pillar.
        """
        w, h = float(width), float(height)
        d = Drawing(w, h)
        ml, mr, mb, mt = 28, 26, 8, 14
        plot_w = w - ml - mr
        plot_h = h - mb - mt

        d.add(String(w/2, h - 6, 'Complétude par section ESRS',
                     fontSize=9, fontName='Helvetica-Bold',
                     fillColor=C_NAVY, textAnchor='middle'))

        sections = ESRS_SECTIONS
        n = len(sections)
        row_h = plot_h / n
        bar_h = row_h * 0.55

        # Determine the maximum entry count across sections to scale relative
        # completeness when ``by_section`` is a ``{code: [entries]}`` mapping.
        max_entries = 1
        if isinstance(by_section, dict):
            for v in by_section.values():
                if isinstance(v, list):
                    max_entries = max(max_entries, len(v))
                elif isinstance(v, dict):
                    max_entries = max(max_entries, int(v.get('count', 0) or 0))

        for idx, sec in enumerate(sections):
            sec_data = (by_section or {}).get(sec['code'])
            if isinstance(sec_data, dict):
                pct = float(sec_data.get('completeness',
                                          sec_data.get('count', 0) / max_entries * 100) or 0)
            elif isinstance(sec_data, list):
                # No explicit completeness — derive from relative entry count
                pct = (len(sec_data) / max_entries) * 100 if max_entries else 0
            else:
                pct = 0
            y = mb + plot_h - (idx + 1) * row_h + (row_h - bar_h) / 2
            # Bg track
            d.add(Rect(ml, y, plot_w, bar_h, fillColor=C_LGRAY, strokeColor=None))
            # Filled portion
            d.add(Rect(ml, y, plot_w * (pct / 100), bar_h,
                       fillColor=sec['color'], strokeColor=None))
            # Section code on the left
            d.add(String(ml - 6, y + bar_h/2 - 2.5, sec['code'],
                         fontSize=8, fontName='Helvetica-Bold',
                         fillColor=sec['color'], textAnchor='end'))
            # Percentage on the right
            d.add(String(ml + plot_w + 4, y + bar_h/2 - 2.5, f'{pct:.0f}%',
                         fontSize=7.5, fontName='Helvetica-Bold',
                         fillColor=C_GRAY, textAnchor='start'))
        return d

    def _make_risk_heatmap(
        self, risks: list, width=130*mm, height=110*mm
    ) -> Drawing:
        """5×5 risk heatmap (Probability × Impact) with risks plotted as dots.

        Background cells are colored by inherent severity:
          - 1×1 to 2×2  → green (faible)
          - 1×3 to 3×3  → yellow (modéré)
          - 4×3 to 5×4  → orange (élevé)
          - 5×5 corner  → red (critique)
        Each risk is plotted at (proba, impact) coordinates; dot radius scales
        with the population at that cell, so clusters are visually obvious.
        """
        w, h = float(width), float(height)
        d = Drawing(w, h)
        ml, mb = 28, 26
        mr, mt = 10, 18
        chart_w = w - ml - mr
        chart_h = h - mb - mt
        cell_w  = chart_w / 5
        cell_h  = chart_h / 5

        # Colour ramp for each (proba, impact) cell — proba is X (1→5), impact Y (1→5)
        def _cell_color(p: int, i: int):
            score = p * i
            if score >= 20: return colors.HexColor('#fee2e2')   # red
            if score >= 12: return colors.HexColor('#ffedd5')   # orange
            if score >= 6:  return colors.HexColor('#fef3c7')   # yellow
            return colors.HexColor('#dcfce7')                   # green

        # Draw the 5×5 grid
        for p in range(5):
            for i in range(5):
                x = ml + p * cell_w
                y = mb + i * cell_h
                d.add(Rect(x, y, cell_w, cell_h,
                           fillColor=_cell_color(p + 1, i + 1),
                           strokeColor=C_MGRAY, strokeWidth=0.4))

        # Axis labels (1..5)
        for i in range(5):
            d.add(String(ml + i * cell_w + cell_w/2, mb - 11, str(i + 1),
                         fontSize=7, fillColor=C_GRAY, textAnchor='middle'))
            d.add(String(ml - 11, mb + i * cell_h + cell_h/2 - 2, str(i + 1),
                         fontSize=7, fillColor=C_GRAY, textAnchor='middle'))
        d.add(String(ml + chart_w/2, 6, 'Probabilité →',
                     fontSize=8, fillColor=C_NAVY, textAnchor='middle'))
        d.add(String(8, mb + chart_h/2, 'Impact ↑',
                     fontSize=8, fillColor=C_NAVY, textAnchor='middle'))
        d.add(String(w/2, h - 8, 'Cartographie des risques (probabilité × impact)',
                     fontSize=9, fontName='Helvetica-Bold',
                     fillColor=C_NAVY, textAnchor='middle'))

        # Count risks per cell so we can scale dot size
        cell_counts: Dict[tuple, int] = {}
        for r in risks[:200]:
            try:
                p = max(1, min(5, int(float(r.get('probability') or 0))))
                i = max(1, min(5, int(float(r.get('impact') or 0))))
            except (TypeError, ValueError):
                continue
            cell_counts[(p, i)] = cell_counts.get((p, i), 0) + 1

        # Plot one dot per cell, sized by population, labeled with count
        for (p, i), count in cell_counts.items():
            cx = ml + (p - 1) * cell_w + cell_w / 2
            cy = mb + (i - 1) * cell_h + cell_h / 2
            radius = min(cell_w, cell_h) * 0.18 + min(count, 10) * 0.8
            severity = p * i
            dot_color = (
                C_RED   if severity >= 20 else
                C_AMBER if severity >= 12 else
                C_NAVY  if severity >= 6  else
                C_GREEN
            )
            d.add(Circle(cx, cy, radius,
                         fillColor=dot_color, strokeColor=colors.white, strokeWidth=1.2))
            if count > 1:
                d.add(String(cx, cy - 2.5, str(count),
                             fontSize=8, fontName='Helvetica-Bold',
                             fillColor=C_WHITE, textAnchor='middle'))
        return d

    def _make_supplier_top_chart(
        self, suppliers: list, width=160*mm, height=90*mm, top_n: int = 8
    ) -> Drawing:
        """Horizontal bar chart of the top N suppliers by ESG global score.

        Each bar is colored by risk bucket and labeled with the supplier name
        and its score. Negative space is intentional — designed to be light.
        """
        # Sort by global_score desc and take top_n with a score
        scored = [s for s in suppliers if s.get('global_score') is not None]
        scored.sort(key=lambda s: float(s.get('global_score') or 0), reverse=True)
        top = scored[:top_n]
        if not top:
            return Drawing(0, 0)

        w, h = float(width), float(height)
        d = Drawing(w, h)
        ml, mb = 70, 18   # leave room for labels on the left
        mr, mt = 26, 20
        chart_w = w - ml - mr
        chart_h = h - mb - mt
        n = len(top)
        bar_h = chart_h / n * 0.72
        gap   = chart_h / n * 0.28

        d.add(String(w/2, h - 8, f'Top {n} fournisseurs — Score ESG global',
                     fontSize=9, fontName='Helvetica-Bold',
                     fillColor=C_NAVY, textAnchor='middle'))

        # X-axis ticks at 0/25/50/75/100
        for tick in (0, 25, 50, 75, 100):
            xt = ml + (tick / 100) * chart_w
            d.add(Line(xt, mb, xt, mb + chart_h,
                       strokeColor=C_MGRAY, strokeWidth=0.3,
                       strokeDashArray=[2, 2] if tick not in (0, 100) else None))
            d.add(String(xt, mb - 8, str(tick),
                         fontSize=6, fillColor=C_GRAY, textAnchor='middle'))

        for idx, s in enumerate(top):
            score = float(s.get('global_score') or 0)
            score = max(0.0, min(100.0, score))
            y = mb + chart_h - (idx + 1) * (bar_h + gap) + gap/2
            bar_w = (score / 100) * chart_w
            color = (
                C_GREEN  if score >= 75 else
                C_BLUE   if score >= 55 else
                C_AMBER  if score >= 35 else
                C_RED
            )
            d.add(Rect(ml, y, bar_w, bar_h,
                       fillColor=color, strokeColor=None))
            # Supplier name on the left
            name = (s.get('name') or '—')[:24]
            d.add(String(ml - 6, y + bar_h/2 - 3, name,
                         fontSize=7.5, fillColor=C_NAVY,
                         textAnchor='end', fontName='Helvetica'))
            # Score on the right of the bar
            d.add(String(ml + bar_w + 4, y + bar_h/2 - 3, f'{score:.0f}',
                         fontSize=8, fontName='Helvetica-Bold',
                         fillColor=color, textAnchor='start'))
        return d

    def _make_dma_page(self, data: Dict[str, Any]) -> list:
        """Double Materiality Analysis page — ESRS 2 IRO-1 / IRO-2.

        Synthesises materiality_issues into:
        1) Intro explaining the DMA process (impact + financial materiality)
        2) The existing quadrant chart (_make_materiality_quadrant)
        3) A summary table of material issues (sorted by combined score)
        4) Methodology note referencing EFRAG guidance
        """
        materiality_issues = data.get('materiality_issues', [])
        if not materiality_issues:
            return []

        story = []
        story += self._section_header('Analyse de double matérialité (DMA)', C_NAVY, size=14)

        # Intro paragraph — ESRS 2 IRO context
        story.append(Paragraph(
            "Conformément à l'ESRS 2 (IRO-1 & IRO-2), l'analyse de double matérialité évalue "
            "chaque enjeu ESG selon deux dimensions : l'<b>impact de l'entreprise</b> sur l'environnement "
            "et la société (<i>matérialité d'impact</i>), et l'<b>impact financier</b> des risques et "
            "opportunités ESG sur l'entreprise (<i>matérialité financière</i>). Un enjeu est considéré "
            "comme matériel dès lors qu'il dépasse le seuil sur au moins l'une des deux dimensions.",
            ParagraphStyle('dma_intro', fontSize=9, textColor=C_NAVY, leading=13, spaceAfter=8)
        ))

        # KPI row: total issues, material count, coverage
        material = [m for m in materiality_issues if m.get('is_material')]
        non_material = len(materiality_issues) - len(material)
        kpi_data = [
            [f'{len(materiality_issues)}', f'{len(material)}', f'{non_material}'],
            ['Enjeux identifiés', 'Matériels', 'Non matériels'],
        ]
        kpi_t = Table(kpi_data, colWidths=[56*mm, 56*mm, 56*mm])
        kpi_t.setStyle(TableStyle([
            ('FONTNAME',      (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE',      (0, 0), (-1, 0), 18),
            ('FONTSIZE',      (0, 1), (-1, 1), 8),
            ('TEXTCOLOR',     (0, 0), (0, 0), C_NAVY),
            ('TEXTCOLOR',     (1, 0), (1, 0), C_GREEN),
            ('TEXTCOLOR',     (2, 0), (2, 0), C_GRAY),
            ('TEXTCOLOR',     (0, 1), (-1, 1), C_GRAY),
            ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
            ('TOPPADDING',    (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))
        story.append(kpi_t)
        story.append(Spacer(1, 6*mm))

        # Quadrant chart (reuses existing helper)
        story.append(self._make_materiality_quadrant(materiality_issues))
        story.append(Spacer(1, 6*mm))

        # Summary table of material issues
        if material:
            story.append(Paragraph(
                '<b>Enjeux matériels identifiés</b>',
                ParagraphStyle('dma_mat_title', fontSize=10, textColor=C_NAVY, spaceAfter=4)
            ))

            # Sort by combined impact (financial + ESG), descending
            sorted_mat = sorted(
                material,
                key=lambda m: float(m.get('financial_impact') or 0) + float(m.get('esg_impact') or 0),
                reverse=True
            )

            header = ['Enjeu ESG', 'Pilier', 'Impact ESG', 'Impact Financier', 'Score combiné']
            rows = [header]
            for m in sorted_mat[:20]:
                fi = float(m.get('financial_impact') or 0)
                ei = float(m.get('esg_impact') or 0)
                pillar = str(m.get('pillar') or m.get('category') or '—').capitalize()
                rows.append([
                    str(m.get('name') or m.get('topic') or '—')[:40],
                    pillar[:15],
                    f'{ei:.0f}/100',
                    f'{fi:.0f}/100',
                    f'{(fi + ei):.0f}',
                ])

            mat_table = Table(rows, colWidths=[52*mm, 24*mm, 24*mm, 28*mm, 24*mm])
            mat_table.setStyle(TableStyle([
                ('BACKGROUND',     (0, 0), (-1, 0), C_NAVY),
                ('TEXTCOLOR',      (0, 0), (-1, 0), colors.white),
                ('FONTNAME',       (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',       (0, 0), (-1, 0), 8),
                ('FONTSIZE',       (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [C_WHITE, C_LGRAY]),
                ('GRID',           (0, 0), (-1, -1), 0.3, C_MGRAY),
                ('ALIGN',          (2, 0), (-1, -1), 'CENTER'),
                ('VALIGN',         (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING',     (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING',  (0, 0), (-1, -1), 4),
                ('LEFTPADDING',    (0, 0), (0, -1), 6),
            ]))
            story.append(mat_table)
        story.append(Spacer(1, 5*mm))

        # Methodology note
        story.append(Paragraph(
            "<i>Méthodologie DMA : les enjeux ont été évalués selon les lignes directrices EFRAG "
            "(ESRS 1, §3.4 & Annexe E). L'impact ESG mesure la sévérité (ampleur × périmètre × remédiabilité) "
            "et l'impact financier mesure l'exposition (probabilité × ampleur financière). Le seuil de matérialité "
            "est fixé à 50/100 sur au moins une dimension. Les résultats conditionnent les sections ESRS "
            "à inclure dans la déclaration de durabilité.</i>",
            ParagraphStyle('dma_meth', fontSize=7.5, textColor=C_GRAY, leading=10)
        ))

        return story

    def _make_carbon_section(self, entries: dict) -> list:
        """Dedicated carbon section: scopes donut, breakdown table and methodology note."""
        scopes = self._aggregate_scopes(entries)
        total = scopes['scope1'] + scopes['scope2'] + scopes['scope3']
        story = []
        story += self._section_header('Bilan carbone GHG · Scopes 1/2/3', C_GREEN)

        if total <= 0:
            story.append(Paragraph(
                "Aucune émission CO₂e calculée à partir des données collectées. "
                "Ajoutez des indicateurs avec le mot-clé « scope 1 / 2 / 3 » et une unité kgCO₂e ou tCO₂e "
                "pour activer cette section.",
                ParagraphStyle('cb_empty', fontSize=9, textColor=C_AMBER, leading=13)
            ))
            return story

        story.append(Paragraph(
            "Décomposition des émissions de gaz à effet de serre selon la méthode GHG Protocol / ADEME. "
            "Le total ci-dessous agrège les indicateurs taggés Scope 1 (émissions directes), "
            "Scope 2 (énergie achetée) et Scope 3 (chaîne de valeur amont/aval).",
            ParagraphStyle('cb_intro', fontSize=9, textColor=C_GRAY, spaceAfter=8, leading=13)
        ))

        donut = self._make_scopes_donut(scopes['scope1'], scopes['scope2'], scopes['scope3'])

        # Right column: breakdown table
        s1, s2, s3 = scopes['scope1'], scopes['scope2'], scopes['scope3']
        rows = [
            ['Périmètre', 'tCO₂e', '% du total'],
            ['Scope 1 — Directes',         f'{s1:,.2f}', f'{(s1/total*100):.0f}%'],
            ['Scope 2 — Énergie achetée',  f'{s2:,.2f}', f'{(s2/total*100):.0f}%'],
            ['Scope 3 — Chaîne de valeur', f'{s3:,.2f}', f'{(s3/total*100):.0f}%'],
            ['TOTAL',                      f'{total:,.2f}', '100%'],
        ]
        breakdown = Table(rows, colWidths=[44*mm, 22*mm, 22*mm])
        breakdown.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, 0), C_NAVY),
            ('TEXTCOLOR',     (0, 0), (-1, 0), colors.white),
            ('FONTNAME',      (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE',      (0, 0), (-1, 0), 9),
            ('FONTSIZE',      (0, 1), (-1, -1), 8.5),
            ('ROWBACKGROUNDS',(0, 1), (-1, -2), [C_WHITE, C_LGRAY]),
            ('BACKGROUND',    (0, -1), (-1, -1), C_GREEN2),
            ('FONTNAME',      (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('GRID',          (0, 0), (-1, -1), 0.3, C_MGRAY),
            ('ALIGN',         (1, 0), (-1, -1), 'CENTER'),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING',   (0, 0), (0, -1), 8),
        ]))

        layout = Table(
            [[donut, breakdown]],
            colWidths=[82*mm, 88*mm],
        )
        layout.setStyle(TableStyle([
            ('VALIGN',       (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING',   (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 0),
            ('LEFTPADDING',  (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(layout)
        story.append(Spacer(1, 8*mm))

        # Intensity & insights
        if total > 0:
            scope3_pct = s3 / total * 100
            if scope3_pct >= 70:
                msg = (f"Le Scope 3 représente {scope3_pct:.0f} % de l'empreinte, ce qui est cohérent avec "
                       "la plupart des entreprises de services. L'action prioritaire est l'engagement "
                       "fournisseurs (CSDDD) et le calcul détaillé des postes en spend-based.")
                story.append(self._make_insight_card(
                    'Empreinte dominée par la chaîne de valeur', msg,
                    C_PURPLE, C_PURPLE2, '◆'))
            elif s1 + s2 > 0 and (s1 + s2) / total > 0.5:
                msg = (f"Les émissions directes (Scope 1) + énergie achetée (Scope 2) représentent "
                       f"{((s1 + s2)/total*100):.0f} % du total. Les leviers court terme : "
                       "électrification, contrats verts (CPPA, garanties d'origine), efficacité énergétique.")
                story.append(self._make_insight_card(
                    'Levier d\'action sur les émissions opérationnelles', msg,
                    C_GREEN, C_GREEN2, '▲'))

        # ── Scope 3 — 15-category detail (GHG Protocol) ──────────────────────
        s3_cats = scopes.get('scope3_categories', {})
        if s3 > 0:
            story.append(Spacer(1, 6*mm))
            story.append(Paragraph(
                '<b>Détail Scope 3 — Catégories GHG Protocol</b>',
                ParagraphStyle('s3_title', fontSize=10, textColor=C_NAVY, spaceAfter=4)
            ))
            s3_rows = [['Cat.', 'Catégorie GHG Protocol', 'tCO₂e', 'Statut']]
            for cat_num, cat_label in self.SCOPE3_CATEGORIES:
                cat_val = s3_cats.get(cat_num, 0.0)
                status = 'Quantifié' if cat_val > 0 else 'Non renseigné'
                s3_rows.append([
                    str(cat_num),
                    cat_label,
                    f'{cat_val:,.2f}' if cat_val > 0 else '—',
                    status,
                ])
            # Unallocated remainder
            allocated = sum(s3_cats.values())
            if allocated < s3 and allocated > 0:
                s3_rows.append(['—', 'Non catégorisé', f'{(s3 - allocated):,.2f}', 'À ventiler'])
            s3_rows.append(['', 'TOTAL SCOPE 3', f'{s3:,.2f}', ''])

            s3_table = Table(s3_rows, colWidths=[10*mm, 70*mm, 26*mm, 26*mm])
            s3_table.setStyle(TableStyle([
                ('BACKGROUND',     (0, 0), (-1, 0), C_PURPLE),
                ('TEXTCOLOR',      (0, 0), (-1, 0), colors.white),
                ('FONTNAME',       (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE',       (0, 0), (-1, 0), 8),
                ('FONTSIZE',       (0, 1), (-1, -1), 7.5),
                ('ROWBACKGROUNDS', (0, 1), (-1, -2), [C_WHITE, C_LGRAY]),
                ('BACKGROUND',     (0, -1), (-1, -1), C_PURPLE2),
                ('FONTNAME',       (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('GRID',           (0, 0), (-1, -1), 0.25, C_MGRAY),
                ('ALIGN',          (0, 0), (0, -1), 'CENTER'),
                ('ALIGN',          (2, 0), (2, -1), 'CENTER'),
                ('ALIGN',          (3, 0), (3, -1), 'CENTER'),
                ('VALIGN',         (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING',     (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING',  (0, 0), (-1, -1), 3),
                ('LEFTPADDING',    (1, 0), (1, -1), 6),
            ]))
            story.append(s3_table)

            # Highlight if most Scope 3 is uncategorized
            if allocated == 0 and s3 > 0:
                story.append(Spacer(1, 3*mm))
                story.append(self._make_insight_card(
                    'Scope 3 à ventiler par catégorie',
                    "Le Scope 3 est renseigné globalement mais aucune catégorie GHG Protocol n'est détaillée. "
                    "Pour la conformité CSRD/ESRS E1, ventilez les émissions par catégorie (1-15) en ajoutant "
                    "le tag « cat X » ou le nom de catégorie dans le libellé de l'indicateur.",
                    C_AMBER, C_AMBER2, '⚠'))

        story.append(Spacer(1, 4*mm))
        story.append(Paragraph(
            "<i>Méthodologie : agrégation automatique des indicateurs taggés « scope 1 / 2 / 3 ». "
            "Conversion kg → t appliquée. Pour un bilan officiel, complétez par les facteurs d'émission "
            "ADEME Base Carbone® et l'évaluation tierce partie (ISO 14064).</i>",
            ParagraphStyle('cb_meth', fontSize=7.5, textColor=C_GRAY, leading=10, spaceBefore=4)
        ))
        return story

    def _build_insights(self, data: Dict[str, Any]) -> list:
        """Generate 2-4 contextual insight cards from the collected data."""
        out = []
        stats = data.get('stats', {})
        score = int(stats.get('score', 0))
        rating = stats.get('rating', 'N/A')
        history = data.get('score_history', []) or []
        by_pillar = stats.get('by_pillar', {}) or {}
        verified = int(stats.get('verified_count', 0))
        total = max(int(stats.get('total_entries', 0)), 1)
        risks = data.get('esg_risks', []) or []
        crit_risks = [r for r in risks if (r.get('risk_score') or 0) >= 16]

        # Score trend
        if len(history) >= 2:
            try:
                latest = float(history[0].get('overall') or 0)
                prev   = float(history[1].get('overall') or 0)
                delta  = latest - prev
                if delta >= 5:
                    out.append(self._make_insight_card(
                        'Progression notable du score',
                        f"Le score global a progressé de +{delta:.1f} points sur la dernière période ({prev:.0f} → {latest:.0f}). "
                        f"Cette dynamique reflète l'amélioration des données collectées et de la couverture ESRS.",
                        C_GREEN, C_GREEN2, '▲'))
                elif delta <= -5:
                    out.append(self._make_insight_card(
                        'Recul du score à surveiller',
                        f"Le score a baissé de {delta:.1f} points ({prev:.0f} → {latest:.0f}). "
                        "Identifier les indicateurs régressifs et planifier des actions correctives.",
                        C_RED, colors.HexColor('#fee2e2'), '▼'))
                else:
                    out.append(self._make_insight_card(
                        'Score stable',
                        f"Le score global reste stable autour de {latest:.0f}/100 (variation {delta:+.1f}). "
                        "Capitaliser sur cette stabilité pour collecter davantage de données qualitatives.",
                        C_BLUE, C_BLUE2, '◆'))
            except Exception:
                pass

        # Pillar imbalance
        env_n = int(by_pillar.get('environmental', 0))
        soc_n = int(by_pillar.get('social', 0))
        gov_n = int(by_pillar.get('governance', 0))
        total_p = env_n + soc_n + gov_n
        if total_p > 0:
            mins = min(env_n, soc_n, gov_n)
            maxs = max(env_n, soc_n, gov_n)
            if maxs >= 3 * max(mins, 1):
                weakest = ['Environnement', 'Social', 'Gouvernance'][[env_n, soc_n, gov_n].index(mins)]
                out.append(self._make_insight_card(
                    f'Pilier sous-documenté : {weakest}',
                    f"La couverture est nettement inférieure sur {weakest} ({mins} indicateurs vs {maxs} sur le pilier dominant). "
                    "Prioriser la collecte sur ce pilier pour équilibrer le rapport et passer la barre des 80 % de complétude ESRS.",
                    C_AMBER, C_AMBER2, '⚠'))

        # Verification rate
        ver_pct = round(verified / total * 100) if total else 0
        if ver_pct < 60 and total >= 10:
            out.append(self._make_insight_card(
                'Taux de vérification à renforcer',
                f"Seulement {ver_pct} % des indicateurs sont vérifiés ({verified}/{total}). "
                "Un audit interne ou tiers est recommandé pour garantir la conformité limitée (limited assurance) exigée par la CSRD.",
                C_AMBER, C_AMBER2, '◉'))
        elif ver_pct >= 80 and total >= 10:
            out.append(self._make_insight_card(
                'Haut niveau de fiabilité',
                f"{ver_pct} % des indicateurs sont vérifiés ({verified}/{total}). "
                "Cette fiabilité prépare l'audit externe (limited assurance CSRD) et rassure les parties prenantes.",
                C_GREEN, C_GREEN2, '✓'))

        # Critical risks
        if crit_risks:
            top = crit_risks[0]
            out.append(self._make_insight_card(
                f'{len(crit_risks)} risque(s) critique(s) identifié(s)',
                f"Risque prioritaire : « {(top.get('title') or '—')[:60]} » "
                f"(catégorie {top.get('category') or '—'}, score {top.get('risk_score') or '—'}/25). "
                "Mettre en place un plan de mitigation et suivre l'indicateur de progression.",
                C_RED, colors.HexColor('#fee2e2'), '⚠'))

        return out
